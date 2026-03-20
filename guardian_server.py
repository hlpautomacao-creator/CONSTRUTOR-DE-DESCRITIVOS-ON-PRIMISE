#!/usr/bin/env python3
"""
Construtor de Descritivo Funcional — Servidor de Geração de Descritivo
Uso: python guardian_server.py  (ou duplo clique em iniciar_servidor.bat)
Requer: python-docx  (pip install python-docx)
"""
import http.server, json, os, sys, re, shutil, io, tempfile, zipfile, datetime, uuid, subprocess
from pathlib import Path

# ── PostgreSQL (instala se necessário) ────────────────────────
def _get_db():
    """Retorna conexão PostgreSQL usando DATABASE_URL do ambiente."""
    db_url = os.environ.get('DATABASE_URL', '')
    if not db_url:
        return None
    try:
        import psycopg2
        conn = psycopg2.connect(db_url, sslmode='disable', connect_timeout=5)
        return conn
    except ImportError:
        try:
            subprocess.run([sys.executable, '-m', 'pip', 'install', 'psycopg2-binary',
                     '--break-system-packages', '-q'], check=True)
            import psycopg2
            conn = psycopg2.connect(db_url, sslmode='disable', connect_timeout=5)
            return conn
        except Exception as e:
            print(f'  [DB] psycopg2 indisponível: {e}')
            return None
    except Exception as e:
        print(f'  [DB] Conexão falhou: {e}')
        return None

def _init_db():
    """Cria tabela projetos se não existir."""
    conn = _get_db()
    if not conn: return
    try:
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS projetos (
                id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                analista    VARCHAR(120),
                cliente     VARCHAR(200),
                cidade      VARCHAR(120),
                revisao     VARCHAR(20),
                doc_date    VARCHAR(20),
                ct_hardware VARCHAR(80),
                ct_cloud    VARCHAR(80),
                filial      VARCHAR(120),
                segmento    VARCHAR(100),
                status      VARCHAR(20) DEFAULT 'ativo',
                criado_em   TIMESTAMP DEFAULT NOW(),
                payload     JSONB NOT NULL
            )
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_proj_analista ON projetos(analista)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_proj_cliente  ON projetos(cliente)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_proj_status   ON projetos(status)")
        # Migration: add segmento column if not exists
        cur.execute("ALTER TABLE projetos ADD COLUMN IF NOT EXISTS segmento VARCHAR(100)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_proj_segmento ON projetos(segmento)")
        conn.commit()
        cur.close()
        conn.close()
        print('  [DB] Tabela projetos OK')
    except Exception as e:
        print(f'  [DB] Init falhou: {e}')
        try: conn.close()
        except: pass

def _salvar_projeto(data: dict) -> str | None:
    """Salva projeto no PostgreSQL. Retorna o ID gerado ou None se falhar."""
    conn = _get_db()
    if not conn: return None
    try:
        # Remover base64 grandes do payload para não inflar o banco
        payload = {k: v for k, v in data.items()
                   if k not in ('clientLogoB64', 'clientImgB64', 'htmlContent')}
        payload['htmlContent_preview'] = (data.get('htmlContent') or '')[:500]
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO projetos
                (analista, cliente, cidade, revisao, doc_date,
                 ct_hardware, ct_cloud, filial, segmento, payload)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING id
        """, (
            data.get('analystName',''),
            data.get('clientName',''),
            data.get('clientCity',''),
            data.get('docRevision','Rev00'),
            data.get('docDate',''),
            data.get('ctHardware',''),
            data.get('ctCloud',''),
            data.get('clientFilial',''),
            data.get('clientSegmento',''),
            json.dumps(payload)
        ))
        proj_id = str(cur.fetchone()[0])
        conn.commit()
        cur.close()
        conn.close()
        print(f'  [DB] Projeto salvo: {proj_id[:8]}...')
        return proj_id
    except Exception as e:
        print(f'  [DB] Salvar falhou: {e}')
        try: conn.close()
        except: pass
        return None

def _listar_projetos(analista='', busca='', segmento='', limite=50) -> list:
    """Lista projetos ativos com filtros opcionais."""
    conn = _get_db()
    if not conn: return []
    try:
        cur = conn.cursor()
        conds = ["status = 'ativo'"]
        params = []
        if analista:
            conds.append('analista ILIKE %s')
            params.append(f'%{analista}%')
        if busca:
            conds.append('(cliente ILIKE %s OR cidade ILIKE %s OR ct_hardware ILIKE %s)')
            params += [f'%{busca}%', f'%{busca}%', f'%{busca}%']
        if segmento:
            conds.append('segmento = %s')
            params.append(segmento)
        where = ' AND '.join(conds)
        cur.execute(f"""
            SELECT id, analista, cliente, cidade, revisao, doc_date,
                   ct_hardware, ct_cloud, filial, segmento, criado_em
            FROM projetos
            WHERE {where}
            ORDER BY criado_em DESC
            LIMIT %s
        """, params + [limite])
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return [{'id': str(r[0]), 'analista': r[1], 'cliente': r[2],
                 'cidade': r[3], 'revisao': r[4], 'doc_date': r[5],
                 'ct_hardware': r[6], 'ct_cloud': r[7], 'filial': r[8],
                 'segmento': r[9] or '',
                 'criado_em': (r[10] - datetime.timedelta(hours=3)).strftime('%d/%m/%Y %H:%M') if r[10] else ''}
                for r in rows]
    except Exception as e:
        print(f'  [DB] Listar falhou: {e}')
        try: conn.close()
        except: pass
        return []

def _carregar_projeto(proj_id: str) -> dict | None:
    """Retorna payload completo de um projeto para clonar."""
    conn = _get_db()
    if not conn: return None
    try:
        cur = conn.cursor()
        cur.execute("SELECT payload FROM projetos WHERE id = %s AND status = 'ativo'",
                    (proj_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return row[0] if row else None
    except Exception as e:
        print(f'  [DB] Carregar falhou: {e}')
        try: conn.close()
        except: pass
        return None

def _excluir_projeto(proj_id: str) -> bool:
    """Soft-delete: marca status = 'excluido'."""
    conn = _get_db()
    if not conn: return False
    try:
        cur = conn.cursor()
        cur.execute("UPDATE projetos SET status='excluido' WHERE id=%s", (proj_id,))
        ok = cur.rowcount > 0
        conn.commit()
        cur.close()
        conn.close()
        return ok
    except Exception as e:
        print(f'  [DB] Excluir falhou: {e}')
        try: conn.close()
        except: pass
        return False

# Instalar python-docx se necessário
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.run([sys.executable,'-m','pip','install','python-docx',
                    '--break-system-packages','-q'])
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

BUILDER_HTML = Path(__file__).parent / 'builder-descritivo.html'
PORT = int(os.environ.get('PORT', 5555))
HERE = Path(__file__).parent


# ══════════ GERADOR DOCX PURO (python-docx, sem LibreOffice) ══════════
"""
build_docx_pure — Gerador python-docx puro do Descritivo Funcional Guardian PRO
Sem LibreOffice. Especificações baseadas nos modelos Toledo de referência.
"""
try:
    import io, base64, re
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError as _docx_err:
    print(f'  [AVISO] python-docx indisponível: {_docx_err}')
    _DOCX_AVAILABLE = False

# ── Logo PRIX (marca Toledo do Brasil) ─────────────────────────────────────
PRIX_LOGO_B64 = 'iVBORw0KGgoAAAANSUhEUgAAAR8AAAB4CAIAAABEs2FMAAABCGlDQ1BJQ0MgUHJvZmlsZQAAeJxjYGA8wQAELAYMDLl5JUVB7k4KEZFRCuwPGBiBEAwSk4sLGHADoKpv1yBqL+viUYcLcKakFicD6Q9ArFIEtBxopAiQLZIOYWuA2EkQtg2IXV5SUAJkB4DYRSFBzkB2CpCtkY7ETkJiJxcUgdT3ANk2uTmlyQh3M/Ck5oUGA2kOIJZhKGYIYnBncAL5H6IkfxEDg8VXBgbmCQixpJkMDNtbGRgkbiHEVBYwMPC3MDBsO48QQ4RJQWJRIliIBYiZ0tIYGD4tZ2DgjWRgEL7AwMAVDQsIHG5TALvNnSEfCNMZchhSgSKeDHkMyQx6QJYRgwGDIYMZAKbWPz9HbOBQAAAtiklEQVR4nO1dSXMbybHOquoVjYUgqZFGst7FjgmH//9fcfhg+2RpRgtJLL1vVe/woZPFxiKSQ1IEojMmOBDQXWvumVUpaICjAtF9mE1no9GobVulVNu2Qojdzyu5XC7zPL/9RghjzPOPdACSP3sAAwxwsjBQ1wADPBcM1DXAAM8FA3UNMMBzwUBdAwzwXDBQ1wADPBcM1DXAAM8FA3UdMWitiahpGqXU4WeISEqJx4wxUg77/hLg/OwBDPAwkEJqo6WQQgillNbaGHMgQCw6QtJaD0T1wjBQ15EBqEgbbYwBXWmt27bdSzlCcBoHU+CQq/EyMFDXkYGhOxQihIDgYg2wB67jbSdJDdT1MjBQ15GBkqrVrZIbC0p0ommf7Grb9kXHN4AFA3UdGUBGGWPqui6Kgi2ufY6NRrfI8WVZNwiuF4PdidUDvFrgDZNCsrwSQtRNvfN51/OapgFNDqT1wjBQ15GBJMGm131gIKafCEejGTqO0zQNEeE4E74UQlDHjIWFdlJIY8wBLIRHW2vNz3iu17Ztq9sDLQjauBAO4zdMI24BPnR7hHiAiBzlNG3Te93uRUlljMHr9uft57e/v52slHAwEpHjOG3bbkuwW8/HPs9+x4gxBR6k3r8UO0XlD9l5b5pmqykM9SiE8PHJLuAHEUkp27ZVQgohmCr2gSChlJJSzmYzKTc6FVxt2KcwDNsO6rouyzIvbk8cOsoxxnAvTB6HeyQiUAiIjYi2aQDfYwzbRAvvhd0XE22PqMSerTSCbLxk1LQ9Ikx7RCT3tUOGtpiOkqrZsw62QSiEAElLKUnvpgoppc3smHqFkvbwjoKoGI6GuraDNtgPYQkishgbSwmA67qu6zqOQ1tIZnchhJBSAhuqqmqaJs/zpmlsJD4suJiKlFRCiKZtILuYGKSQnuc1TcM/2e86juO6rlLKcRxjjOM4UkoYTk3TpGnKpOX7vud5PPEDAyqKoqoqTFBK6XleEAR1XTOXYWkmhBAHsRfrVpZlWZUYcLtHZmJ3HMcZj8dKqbIs27Z1HGdf+0opY0zbtoh6w0nTtq0X+FmWFUWB3n8w2VcGR6MZYvuh5EgpLy8vp9Np27bgteBzSZIsFgu27z3XC8PQ8zywTyBTWZbcGhMkkIbzHgCO44Asy7LM83xbhdsJICQp5Hw+j6LIGON5HlrG4DHUOI4XiwW36SjHdV3f98ECoP0CO5GTYYyBh5Cp0fO8KIrQJu1n6kJJrXVd12AlWmvXdcfjcVVV23GwA9QFTYHpk2Iqq/KAOgoiCYIgCAJmFkqpfe33+CN/9sMA3lF7grZ18JrhaKjLBiml67pBEIC6kGgnpSyKAqTlOm4QBJ7ngbSg78Fsg/giImw23aUuALoAJhHRZDKZTCZ5nq/X6wP2EgO0wSAIoijCvRfGGAySsysYmRzleJ4XhiGwEJRQVZXjOJBXZGUGItLV6tYWv5A8++7VcBzV06zKsizL0nVdEHBvBfZhPzpi6grDUGu9z1EJwB5h7swj9gkfiGghhOu66AvEWZYlWANZ2uax6IfHRF2MQE3TrFaruq7rulZCNk0DhSpJEjwwnU5d1wXSAIdYKNV1zfkNLA/xGL6HoGOEA49USgVBoLXOsqxu6gOkBdURxluapmmaoncQM/5ChCqlhBbz+RwoC8LGk1CNeJwQpG3buo7bM+6ZHeyDuq57OnNZluv1+vLy0vZ23FLgftlCHY2BuzmOc5i6oDg0TQOdcEOf++kCtGfrDo7jfP7j97quqVNrWQIf6Pf1wDFRFzMw7K7neY7jkDYQR9gY1+m+73YL7wKngaz4YDdIHY3xW0yTrutCmEgpwzB0HCdN06Is9g6SNv4JtuJg+KEj1tCEEJCrvu9DrjIZQzPEk7a+hFHZzBsDRvv7NCVNhu0rMBEiKsuyKAqgL3V212YF9ngdMGAeg+M40B3Kutr5/Hw+932fSYK6zBK5R8ayCg2V0vd96M8grc1cjoSoGI6GuoBAbC8B+4UQdVPDOGmaBva667pSStjxZGXiwTfAiE6dpmHLKABQjZtl1MQwoigKgmC5Wu4cJ5TGjcPQGGiGRAQvArpTSnmeBwYBa8o+HgJikx3Q3YwnYfmmWdwdMPelo7Z/Ncas1+swDMMwhNjEOB3H2ecLhcEGZRXjH41GUsry5nrn87gNDpotz0JKqZvd7TdNEwSBlBILpZRK0/Tq+oq6NBTb9UJHohweDXUB+1mjwxIDC5m7Myfu5YzjAWh3eZ5j10FstuHhOA6QHmqM/SvawTPYe9v0sh2JIC3IrqZpYN7wePCX9R8mXepIyJ4yRwuY9qwo0O0iQGncZ3fRVqoh1rCqKsgWtANrB5r2zkZgExKR67pEBOblOM54FOV5bgf0lFRhGLZ1Q1suadPqfeOEDimECIKAiOI4Zj2/N/6joCvA0VDXQ8E2q+C5hkc7jmOmHDswqo2mkmQmQWPwd2PL+XkWg1rrKIqKooBjmu7hqd8HNiWjC8gQcH2IHdsL8oSQ5zmwmSlHKbVPM9wH4/G4ruuqrjhC5bou3LkPHc9GdZSyaRo4Mx7awmuDU6YudlKVZam1zvM8y7MdT9KtS00bXdVVVVdVVQVB4DgO/H7mblgMqhHar+oKCRmCDiVM7AOOO7FjAyILwgq2EIvQ+0Sx7w9gN2wWbmIGD+QRrutuAvG6Nca4jhtFked59u2/9wHW+rTWVVUVRfE4bvWq4GSpC3YCERVFkaZp0zT2btmZDRx3NlYCDmgMGR6ii5WxmwQ+Eg7LGjIIcz1uqLavDBTLVqKdp2snav15YHOrLEv4UQgq6wPbqapqPB4LIWCITiaT0WiEoOKDAJ4SIirLMssyzPTRGsErgZOlLo4dcRBMdjlT26RFXbYR1Bvqsj0MmcVi4bruaDSyEyPY9+j7fl3XyJl6HB6gBU6hZG+n53kQa0B9xLXruk6z9E+uzGbinTQGdSEw3TTNQ3lE27a+74dhiHQKz/OICPG6hw5JKdU0TZZlcMk+mlu9HjhZ6vJ9H1ZWVVfUbdXmcFRnbpG1hexJN8bYPzVtA38aXJHs94OjAnHVtm25l4fSGDzvICHkQ1RVVVUVonmIPYC6qqraqdk+DkyXNQLxBTvzQB7gPgDNE9HZ2Rl18bQDt+jsA8y9KAqWe+yvOl44Weqq6zrPcxvp9/maNp5uuk1+71GIIVMUBQgJKRT4Hv59BLKhHz5unCAtRMPyPM/znKO0rW7xGYrrI9o/ANxgVVVJkoxGI9/39QOVT0dKBOghtcqyJGPgu9/5/D5ygQs+SRJWzjepz8fjIdyGk6Wu5XIJhc1RjhCiampB5Chnkw9Bt4neh7HWdVwkZ+R5DjECM8mOBDCrfgSBwQGNdGG2N6i7+4kfa3XrKIf0jiz7xwEHoCHB4D8EhTwIOBgAHz2HtnY+fGDoSGRDPOOHWSDHAkd/BResBUSEqqbWZFzfM4KyIjdEhqhum6qpSWw+t0ZrMvjpPv/VTe0oR5DIi3y9XsON4Xke7A1Q2mw2g9vwwDgR++IYFyfFGkFlXa2TeJ3EjW6FlEJKQ9Qa3R9J29z5Utz576GA3m0VDnlemBQrw9RF5/a1o4mMEPyfJtJErTH40PsPrTme2xpdNbV0lCbT6DaMRlc3163RJIQmU7eNkJKE2JeAfyxwCrJLdHmxoDGo73efsP4+XNHgyBjwD/4S9nDAEguDsN/p1iDtODgnbRhBVVXB2LCfefmYKQ6J+K4HVwqnWfB8n6YbKeq6hiZcVRVi93EcIxfMzpl6mu5+Khy97KIOv6GtQZRl2ZNZ/2RlSLRtG8cxBBekEHXmeBRFdNDNtTkd2FEOswM4FWxCpf0XPD052DSMXA1kSCDGbee2P1WPLL1hwSKgt1wu67o+ikMlD4Kjpy747jgVED6rJwzz2wSjtU6zlDMV2W5pmiYMQzpod/VkHbPnLMtApfTit3ly+I4zmIkoL3JYUPBhYqhPSO12ViecH5to5N14PafaHDWcDnUhGot89ids31ZRsN+co8Rfaus47YF27FgZXq/rGplZwjoySC+bDM6kzjSW53lVVbC79OGLfh/VHdROGK5FUSyXS/uB09AJAUdPXQyc1/u01LVpnIQUmyyhPM/trHnoeG3bwvt3n0GyNMjzHMhtU+bLoJctr24TL42hLluCzxM8LaljshxqQ8IUM0d7SE/Y6c+Co6cuaC9QLWDGPK363vOASyGTJLHlD9ClrmuYXvugZ1DhQ5qm+AkWDj/5AnbX3ugfCUMGYTce1dPiOjJ9hRBxHMMVZOuBp6ETAk6Eum6vHOq+fNpe2KByXRehmB513Ud2MW823R0bTdvA690b9iNyHR4K2wohhie7KwaKogB1PW3OBOQ88kLiOKbuvhraoisezPHCcY+eOraKrDalFDbsCeE2VcoYKWRZlVJIXPkCCsHJKMaVAwB/JgJfELNSSHbAGOsk9YsdvjDWnXNAbq01AndFWaRpimwvYx3fpPuJMg4rs9CW3dlwnBm9vr421t1B+8b2RBP9OXD08S47fPRim8FowVe0g8B6CRY28CkPlrRleehOpZ8FdspYXddZluEA2EPbsY+lchBCa+0FPoIQp+d/34ajl10M7CKnJ7UTbhOmrMwpWPy2TgUPm+/7+9rhizoQ6sEZzcPpHT8dmrZJ07Qsy0doaLagY/HFN/akacrxjKcd86uCV727DwJcyfAcLdvKIRHhyieyGHPP37UT+B4B2d1jgyOYzzHgPwN8PgAf6qaGAfbQdvjcJ8fQkfFcVVWe53zNE52Ke3AnHL1mCMD+8Y2zz71h8EwCe6R159GBYICtGeJkZJ7nB+6W+lkguivscak1EeFEyXg8flA7TFT4J/t1kyzllLFeDPBJ5/Eq4ERkFzSuR7DYxwGroHxonx5iheNmpdd5b8R2vog2+qHH+Mm6xQ33/OA8P84BmK1KYqeqH56I7KIuI+ll+uIEOU5vFd3tAPsO59vmh7CuWHttsG1n0qMCytsORmHdbdjTCY/dN7gPToe6yPJTPbeaYbo0XL11g+cB4IEZ6zq3VwhcvQWfHxd0QmY964e4JYGIVvFayjtlTU4YTkEzRAgly7IPv74nYxyptqpwW/89ov2uuA44rhRyMpnwtdXAHt/34QY83BSfWUSu3Su8OgIJX7iWkIgMGdd137x5Y9PDfRQ55nQc5gKx/fr2nWm10dpzXEEkiCSdqF54GtRF7M17ziv8tw0SVm84inUCMRz74D0RuY4bhuETKrE4bKqkevRNJEcEp0BdNt7jeoxn7Wjb+c6M+QSoiyClO6E6Go1Qx+SpGjfGTCYTDgxujvO8Phn+JHAi1MXJE/atMk/ci3UVh30Ayc4ePAFbwja6fM/na4yfrH0hhBBBEKCj0+BH++AUqIusUgye5z23psH3rtFWQPlZ+30ZkHKTzOW53ng8xiX1T8iwpJR5nnued35+zod6TlU5PHrqYpxm6nrGvui2og9tGXinIbt4CijYR90hgKdqXwhRlqUQYjweQz88VbWQToO6bNHxiCtgH9QXWRVJbGG18T4fPw+G4Ar8AKSFeMPTejVw5FlrHUWR67gnsGj74BSoy/7nc4doEf/phYDYZ/isXb8YeK43Go2gEyJp6wmVXq217/sonYECYtSVcj89OPpZIW0CSEBEUsrJeIKftm2hR9AeZ7XidAmuXAfrhQdFa42ir49IF/q5YC+OvVbj8Xg8HiNdy/d93/cP5E/y5VZ8DOeHR4EQIcRRt/l8fj4/h8BkvYOrRgynJ38y9Bx3QgjXdV3HpbuGEO5yeExRKTK3Bb6IcNiJj3WRVXP11eZe7AR2zNBdF2g0ilAtgUXWYWOSL8lgx5K4ewnpThDdkWdcSxr4Gy0Uv94/Y/OVw4lQFwOoa/uKiz+j2ziOw260MAyZN+uumN3mOODTlf95SbCvr5FSjsfjbeo6sHp8zAT/xOfDz3O+JRLnlVLT6RQnoKm7H9L8pCtTnxaOnrpswNai1De+YRuJr7V5RLO2Gw2VWvl70V2ctjn8cjzuLxtxWanGBKmbGijnMHWZ7u4All18lcBOsK9zZFnned5kMun1YkvXI4Wjpy57S3hrlVKj0YgsvMEDj/MobjIDlYOsBaiFzFlhoL9Yev5TwfYxEN/3UWiTTSm6hywy1tU3dlbhvueFdZs3d9G27Xg85gAAP/l00/05cPTU1QPdwWQy2S4y8IgjVYh4KqmiKLLrhvADTdMURYHak8foXAY2u64bBIHruhDy9w+Og5fZXo3DN2OzDmmTN/RDlCDcJvvjhaOnLtu7YGckKaUmkwmCy1wF5xEbZoxRUiHdjtPbmfsSUV3XuIPtWUNtzwFMP47jALP5yCPrhFCtD/gY7qxDd57/MHWxj9G+jqosyyAIYNbiyROIcBw9ddn5fmyggx1OJpPLy0sQGCPNg9sno5RCXTxQaa/0m31O+WmmdA/YFHUxm7/iUVyeycBxHJRgR/0EYV27bad97QRbPeZyE/ehLrLOL0spy7xwHAcVWMQmIfvo4eipC5y1rmtb5TDGCENt3UgSk2iME19Gaxwl2vefJOEqBw/zf/Pz88lsagTVbaNcR5Mp60ooiSJUq3jNhcKKqnyBk0rdUEmY27/7/iPrGDXdjVuQdRXkfD5XShVF4fs+DEphjCOlEoK0Jq3xGW0qIfBZdDpCo9usyNM8S/MsL4vWaOU6t4XFpCApNBmN6u9KCWNIayWE27nzlRCB52VJGvr+r2/fecoRRKpbTb5M6rnX9snhyJSZbUBUF2wYNYhx+VlRFPgchqHv+3meJ0mC0ob8rm1dGGO00U27cU5IIaMoCsNQuZsy4T3nWNu2qCZuH4h+GY4rt/4eethyM9giBV86jjOdTqlbxqZpfjh+uykhhOt7WZYVRYF1yPMcAfeyLHvd4X/76jK3besqhdt/J5PJarViU9bWFB6t4f8UOAXqAgdkTUZ0dynzcUbYFcjOtu/oNKbvhpBC4qIV3/eDIPB9v2pq0wF1wVMhRFEUqB1OVmz0ZXZdd3/xn+j+7oSeXmebqcaYIAjG43FVVUj/K8vSkWqnjQWRx0shOjDGoGYKHquqqiiK0WjEZABV8DZ0tqe+a6u153koxB5NJmVdJ2ki9tQmPxafx9FTF1leY9/3EQktigJJFXz7Em5sDsMQxQdAdezsQh4TrvvENdRsqWuj2b7nhA+UOYao/Nmz/wEwX6AuEYxvBMFdFzw7OugEtyNUZIl9rrTA7vWiKOI4DoJgm+MYY/aFBHsCFtKvbmpcuMC+E+79Kdbm2eHoqYsPg9xyRyFc162qClfecuYBsAo5gcAtlBXF5qEcia0o4jNjJA6eoLAlrg1jvDHm9v5QemX0Znt9wBr4m/Pzc9d1ocvhrtUfUhc+s6sDxThZN8av+BJrK+8WKTfG7DNNsbZgiHmeB0Ewn8+/ff/GC2sLsUF2vRCgdhvwJssy6ISsmfRkFO8feCq+IdqUbKOtqgVEJEiyvxH4ZNcgff1mgG1b2pJkPB5PJhPMBfV+kPJ3uDWbUCHAmbSwJkx1UB/Q8n3WhwWUEKKqKnjnfc8vqpK2Mg8H2fVCwEoOQekvC5ScG41GursOjTMJ+PJ3jn6ydsdJ2Rzt2STmGs25VFxeYKefgF69H5lH6/u+bW6BHu5/gEB0FXQh8KknnYiEEEmSwIJlSrDpfOfYkGFIRAhqCyGm06lIEziobOF5LDm+R09dnufBpEa2QVEW2uj1em2McV2Xo8lcALaqKiY26rwUZKW6U6eHADPajrqATHxLM2SgzcuPZctd10XseLlcwt9TliUReZ7XNM2+bJMeYWwYE93m2rLdRVaYxNzNURRC7Gsfujd2JwgCXF5/dnZGUhhjsOz7mNqrhaOnLtta2CgnQmqjV+uVFHI0GiEQDL6IaCkR4QPXBeYAMVMaca0T1ymKIssyDpXafQFsGwy9e65HRJ7n4S0kBwFA8FVVua7LAYAHgBAbJ4QUQknXUVprJffe1SOE4Lumich13el0Op1O8zxHKia0ZeKMlj3SBZgtlYK/dL1eb9LKdt2ny955yJ8gCCD5DwSmBQlNxvFcIqqaWjrKc1RW5BcXFwhVs/g6FtKiE6CuA4AL0Jum8TzP8zyIMtMVEaeOQqDoQ6Bx1BImWdu2i9WSfYxkJWQcllRN06RpiuQgVrqYgGG0PO4eedHloUNis7dmn14Hxs+j5VLFD80zYt865P99Bm+MqaoqTVNE0jh2sm9eO7+/vr5OkmRb8zwKGjtZ6oIMaXXblm1RFq7j4rpc3/chrIgIZedd182yjM3xqqpQ2hyRMXsP5d16igdAG50kCYssuntCcdPIo1J+uYWyLEG6YBb7qJ1/YvU1yzLOCbw/4CA2Ox57je/rl81adi899DABZ07bMFDXzwc7ab1uNrm2oDrf8yGyxuNxGIar1Yq6Ana3GVV0e8iPrEvdDu8rQjqGjDZaN09viRlLE+u5N3cCT8fOJHqE2NxJFQe8FPyT/eIjzun0tAbTpTU+tJ2fAidLXUjIwKF9smQFvq+qCv+M4xiext7ruEiDiForfZGtsgPYvLkpoCNsNLKvgqsUsv0TxV1vjb39MsTssovo4U5tbscOQhzOnd/5/UM1Q2Zn5uAhztcJJ0tdAK4CzMUUYQNAz2l1y3iJogSbt4zRZkdaU0+IHeiU77rp/XRbxtIi9QeBHe/mLw8kYW372f6Mh7M39/vIzG1n487nf0iNg1fjFUHvLKOtIvJfImp1i/sw7FsxlFSO3BziInl7Zuz+7mAmMNAPVyt/quOV23G2fSnkO6nowPP7gI+WPBS5mSUd7veeVHdEpEUnTF3beAx03yQW4oov5cAnLkhscjiMpruUZh54hI8Jqcdln+rkcg+9uJfDssgO6dqc4tHwQyu0J2P3Kaj37+tYbC0bjp66ODlic6AIH3bt+EYZsyin7sJNhox+ROhpF7AdZRtUf8a4Ogz3JJL76HIP7feH3p0n6eg+fb1aOL4TaQMMcCwwUNcAAzwXDNQ1wADPBQN1DTDAc8FAXQMM8FwwUNcAAzwXDNQ1wADPBUdJXXz8kQOadqKdftJqbgMM8Gg4Guqyb9LFkUdOaTfWZRiiK776M8c6wABEdES5GnbAHmf4+JIMTgViaXaMWTMDnB4cE3XhA2ed4WQe3z/Dd9TQseV6DnCqcDT2CSwrO9UaZRPsay5xNygdz9nVAQZ4FcAnF9mmGlwXA7xyOD4ExVER0RUsxpfHm0Y9wAA/H/jOo8OyiwslDzDAAAMMMMAAAwwwwAADDDDAAAMMMMAAAwwwwAADDDDAAAMM8GcAFU0Z7EI1qDu881c7I4nulr0iK+DLBXv4MX4GVUhs4KqQdi92ayiCvD0ebsp+l6eGelm9odpBZy7Pw4dchFUxGQmNdi/8AY/1/mnfy9lbDftDrwo4SqHzID3PU0r15muvHu9Lb4Mwl+1qQzwLe150NyK/c8DbIfveIvdQwl4EXvl96NHr10aJ3ruAMAx7I0FTdolDIvJ9n4iCIMAz3KxdCNOukIhmGZfwJV7nxnn1dq6JvQuo9b4ZDB51HMfzvG1y4g7ukwCx7xlsOc8fxdp4iKjxs92Oveg9Ctk5QyGE67roCA3yJLk8Nt1lH/tmsRPnuIseffLUemS23UJva7e7VkptMx0ms32j2jlO/ry9tr7v79ypfaOijqL2jZlPstpfKqV412zCZr48nU5tRsC7tn1Cz/d9xnVuhw5uJba+l99jsxUMSQiBQfYoapshAvAwho0pM5Wi9Cm/Ana26WM0GqH+IqpXYXwo+Ok4Tu9sFX9GN8YYHGS0bwm3X7HPifBpEZ4tesFAuVwdxod/hmEYhmEURZ8/fy7LkqnUGANuhHo5qA2FTu0KOuida9vYR7/sauX452g0chwnjmPU7QUtlWWJxRmNRtPp9Obmpq5rLmO5Kdm4675b/p5RqleKCgu4KdBsVaB1HKdX+Yo3G3PZ3NetNbIusQvMuYMgqOs6TVPUULd7FEJUVWWsei58Iu7AmW78NB6PXdctyxKVIHnAdJeDBEEQBAFYal3XYRgaY25ubqqqwsN49927dzc3N0KIxWKBjeO65lylEnt6eXnpuu5qtVosFjgq0TtnZOMkPv/222+fP3/+8OHDP//5T54X1xCzkfPs7KwoijAM0WMURf/973/RC68qTmPg0BNZd3fb+17X9V/+8hel1Hg8/te//gV5uNGdQF14AaWmeVcYL9EZl1rDcvQKnNpzxhBBpXVdg0p7Z0NsjOSmuCAAMNL3/SiK5vP5//73P7JqYdjPkFUHBFVOMAXgrs0v7GsCeuMHdfExFmMM6jtiEYxV4Yap0Z4RvpRdUXN7nIwx9gR5AKIr1srHQJlibbCPYPcwhvdIdKXrelOzp2w3aKySx7TrEJ2972DPZVlul+GyNx2Crmka0BiQcrug3sXFBZZ3sVjYy2LPyBjjeZ6tufH0aQvTGP08z5tOp8vl8vz8nLkAMEFKCcTmZW+a5vz8XEqJaoYXFxeLxeLm5oas0qSoHuy6Lt61mbXpStqj6DO4f13XZ2dnUkoHEgAraGvndoUY1LF2HAdsD8vXtm2apugmCIIwDDHiqqryPFdKBUHQtu3Z2dmXL1+IaDQaoZemaZIkgdhxHGc8Hiulrq+vMQfecpRjxYq0bVsURV3XruuOx+O2bR3HSZIERSJ5jx3HOTs7C4IA5RXjOFZKoeSkMWY6nX758mU0GkVRBIFZVRXKXQMPRqOR7/tgLufn5yiXTERRFGFGECm+76M4ahRFGG1d16gR7Pt+GIaQexAdNgpGUQQWRkR4hYkQ6xOGIerb4yd8CXUXNSOrqprNZkmSoORxURQYEkSElDLPc1SpxQSDIIBFJ4TI8xxfRlHk+35VVWiTiMIw9DwP3DqOY4hN3oUoioQQTdNgi7FimCleKcuSi8e2bYsa00qpMAzLskTlQWA8Nq6qqiRJQHjgZUqps7MzvIUWwJellPP53HXdOI4x+DAMsc5FUbRtiy95nWHdTKfTLMuAZhBHQHJoH9++fcPDQRDkeQ62LqX89u3bhw8frq+vP3z4sFgsRqORlDIIAq319fU1Eb19+3a5XAIn4zg2xozHY6zSdDr9+vXrarWqqgorWRSF7/sb2QWqYNML+IGpuq57eXkZhiG2fDQagRNAsY7j2HXd2WyGlXIcpyxLVEydTqdaaxQsllJGUWTzda31fD4HimMRl8tlVVXMyGezGfCYxbQQ4uzsbDQaoaK2ECLLMiA6683QXsA44zhu23Y+n2OlAJPJZDQagcNhe8IwnM1m4AWsbn38+HG5XH769AmUtl6vjTFQrJVSWZaFYTiZTLBQEBRKqel0GoYhflVKLRYL5heO44B+mF8IIeI4xjoDd/HWaDRqmubz58+TySQIAmZnYIrv3r1brVYgDyklHr65uZnNZjCesfhY2Ol0CqwCMSyXSyAf6FZKuVwusywLgiCKIlSIVUqtVivWS6WUQGgsC4pQY7+wDlEUgc3xK2CgeJ5Nj7OzsyiK4jgGJSdJAoKE+ndxcREEAabTti2quWOcaA2rNJvNZrMZlnQ8HqM4OqgdBI8K6CjVC1EDulVKFUWBFtI0BSZDYgM/gXgQSkVRYMD4Cfvred75+XmaphgVFgQI3LbtbDbL83w2m0kp37x5A41Xay2JCAw1z3OMBlKF9RAwhjiOV6sVhECSJIvFAoMA566qKo5jSAzIE6wIC9zxeIzPvBYYMRHhS8/zgBNAX9/38Sv2ybYysyyD5QNegHUHBmNBl8slpNZsNqPOvtRaJ0kyHo/BsG9ubsqyZCcSml2v1+B5IAC+SgCUk+c5BEuWZW3bQofEZ4hZ9kmAtiFeRHdoGn2laXpzcwPO17PWmqbBGCCNJ5PJZDIBt8KiYb5QIIFYSZIAq0Cupru3h704eBj7AlwJgkApxe8C+5umKcsyjuObmxtWUHnBgV43NzegImx6nudJkoA7MIDngotxjVxYYr7vN02zWCyAYzDJgLuO40wmkziO0T4r2+j669evEHRYZ0jRq6srYIjtr4KG5fv+zc0NCmHjy7ZtoRRgkJeXl5gg1hk9VlX1/v37oihGo9GnT5/ggIAFvlgswNqwaBAwYOLwt7Vtu1qtfv311/Pzc+jARVFg7yRvLTQHW6PF0gMtsMHv378PggB4ANUCtL5er0FLruvafk8iurm5McaA18IwzfMclPzly5erqysInDRN4ziW1iUZ0B+gErCmC/LGQkdRFEURWUos9iyOY7Q/nU6hI2VZtlgsrq6uxuMxRCgYB3YdnQJ7gArQ3GArEhEGUNd1URSgQ5g9eZ6vViusOFuheZ6nacoGIVPR5eUlFAnguu/72ABebZiX2NFPnz6t12uoeUVRYKnhooAZUxQF4zrwDLMoyxJzh28Q67Ner29ubjBUbH+apsvlsmkayDSQHNBivV7zhSVSSnANbIfoPGxgT1mW4S2W+fZFXVBYwJrZI41Fo+4ILISh7SbBk6A6ZvGO4wDrAGBq0D+Zp1BXSBqEylIRKzObzSDl8jyfTqfYFBhR8FpFUQQFFfP1PK+u69FoVBTFt2/fFovF9fV1kiSTySTLMnbYwKGwWq0+f/7873//G/x6MplQ50eQ8IZBZQKZQROArgieAS6SpinUUN/3ISiwuHDowTTEMoHWfd+/uLgwxkwmE2wz6+WY/Gw2++WXXzzPY28SbDzIQzhqjDEXFxdgurzWl5eXUGniOCbLJ4lfwTng6GcmiifBO8MwhJRr23Y6nc5mM9PBfD4XXYgCRAUiBJ+DXsc2CVaJHZUgNiHEmzdvrq+v//jjD+p8LcBU/PPy8hI67cXFBVmwXC4hXVlqQQK/f/9+PB6DriBMPM8DsxNCJEmSJAnUIeA91gpLCnYAR99yubTdRcCqLMuKovj69avW+uLi4u3btzayQp+HU4E3F5L84uJiMplcXFz0AiocvDGWv4qIwN2xVlLKqqrAxdgScRwHbqckSaDBNk0Dm221WoFcoTvA5YDPrI4ynUBYYemgSUFaYgtg+FHnIwGOjcdjCEljzNevX//+978DvYMgWC6XeHc2m3med3l5CT6otYaxkCTJx48f3717B1EMnRZr5bquozqAkAX9GMs/Jqw4HXi27OJRUNPDMByNRqvVCjaMrfBgF8FybNMIpOW6bpqmVVXBwADTZbcpMwBozPhGCLFcLuFj0FrDA0F39StIG1Zcwb3w7ng8JqI0Teu6hqIPUcDjxE9EBE0VVi/WGuILjA1MB7iiu2sImJZg+czn8ziOQQbg69jL1WqF1V+v1xgt8BXsCXYz7PLpdDqZTJIk+fr1KxFB1cTWqO62Oei9YHPG8pKrLtQDRRfD5u4wcfDpqqouLi6g+sJ0ZD8+BAvkFQxFABPwYrG4vLzsCRCmNDACIInrutBF7YHBpoBRSkRXV1eMu6C39+/f/+c//2GFE7NA19IKwanuMgjHceBe4kQCjFZ3ABuBEQbIBl2Ax4whATFgkv31r38tiuLy8vLTp08QLTC6YM4lSXJ+fg7t7/r6+uzsjDqXu8MW0Xq95rXgYBdWE2wMXAH8gDpvBC8lj48dyhy1AHMFq4AaA8kG3tN20Z6elxJbCNQnKwyKzcav0LKwWGx9TSYTMEjYzUyZmIXpoltwfiRJIqWExogps9rgeR6okSUbfHRYO6hJ8Olh0eAmhVoCvZyno5T69u2b7MK1Z2dnQojFYsHyGUyK0RczBT1XVeW6LvyNMPNA2Fgu+KbwMCtdzJswR2BkFEW8R1hD+BvhP1itVmmaTiaT8Xic5zkLZG4fOpvWGp49tB/HMaQ9RA1LdRAPCBv66mg0gq0PHIOiyAIEjkQYTr/88gtMD+rsVc/zTBdTRQtQetkBQ53zg5ENggsGBb5EO3A2ii5uAcIGtUDfgc7CmC/uBlSWy+VsNru+voa/9N27d7///jsY5du3b8FEwLWx1w78ldDWYHfB9/L9+3cmcQyRGYPrulhxTBUi4uzsDM9gmbB24C7AV3i34DIC8w6CYD6f4y0gru3ChjLAfnz8hdYHU7JtW6jjPE48dnZ2VlUVSx5wQXD31Wr17t07aPaj0QjGFXwSrutCCEP/+f79O+xX3/dFl4oRBAF6B5YHQcB8HUZwGIbT6TTPc6gAfP0bx98gsc/Ozpqm+f79O2gVkoT9kNAbwQshnd68eQOkZ1wET2VsxjrjAewIawGQhFhbWEqg1fF4DPYEmx6zg+NUd7EjjBnmxF/+8hdMnCUJJstCgwkML3JGCAQL0BqaFQie/cCwQYgIemaWZVAFgTBv3rwB32GSBvEAg40V/mKdC6YEe7axVnBBa62XyyUvGpYIU8AqgWXAsIemV9f11dUVeNObN2+CIJhOp3/88YfjOF+/fv3w4QOc7UmSzOdzuM2aptk4bJiDYiNhQzOTA8/AZsPzi+U2XaQVdjxGA2EKjd8Ys1qtWIhxYAGGIIIV0KDQKZykovNfF0UBF2XTNPCwAYPhQoBbsizLNE3Z/LXJkhUbcF/dxV6LoojjmE1zzIuIEFcAzeOb6+trMBEwYEyKHc1KKYyEdUvwWuaLRVFwU+A4RIT4JjYYa2tHJ8GGgZGIlUF3V0qVHRARhAy79dE1S1ToFHigbdv1eq06t/hqtWLyA6tK0xQ2+pcvX6DPYKnREVNXlmXg+rDTsMLr9dr3faUUi1OygDcdjgdomHmew2LET1dXV2gcNuF0OgUKwTXHEqCu6/V6TURoCl5Nx3GwQWx/ms4NaLpsgSzLQNII22it4fXl4DVQgnEbD2utYdHFcfz7779jRxaLhe/7379/n0wmZVleXV0ZY/7444+PHz8mSRJF0Wq1urm5SdOUGRNGsulGbiWesVgkS1femYXID7PtzqC20j3FVvqvvJveJi3nOzuaqDPh7H/yk/zNxcXFu3fvYJ/gAWa09oy2x0ldOh965G/EVsIElFseIVs+vcHzizBl+XVOU4SEUXdzjqFMgqLsL+016dnDvZVkdVR1mXjg5R8/fkSghu4mPfNfaJ60Z4s59N9bMepSq3vrKTrPEH9jmwY8DNuf/iDgrslKwmYM5HVAX7xovMLbb4kuFVbsSTjG+PEXjYBJ9RIaiQgOvI3o3odDNnVtg+ySF5WVl8mqqmMld9vt24N2diV3bxPb9qjsJVZ30zTfvn37f//3f9vj5M/bjdvkZA+mN9qdq2Q/z7/aKrH9PWPVdke8ets/9cCmXuxuj1v1KAcRaiwL9KXtXpiofojojKO22+0ASCtfljdrJx+njt3gG15t+2HmiT0k6aUpb7+4cyV7g+89w5gGNYetA7LYlv0WzG97qJufbL7IXdroxRzOTlTvPdmDnnDgz4zHNkLvbKonBMDOeST8eo/HvH///rfffgOh9tLqe6+orZRznr60jpz0qNEeNr9i8459ohtLz7/ybuEBe6jbspHuov723Htd2Dzo119//cc//vG3v/3t7du3vV6Y8dPW4Robj+3l5XRyG/NstWVbobAXweYv9hxtebKT0WwzYnvM/Ct0FhZrsktS7S2RTeS8fezmOdAvWfyCX+wxTdY+8M//B1113cdJk51TAAAAAElFTkSuQmCC'

# ── Cores Toledo ──────────────────────────────────────────────────────────────
if _DOCX_AVAILABLE:
    C_AZUL_ESCURO = RGBColor(0x1A,0x3A,0x6B)
    C_AZUL_MED    = RGBColor(0x2E,0x75,0xB6)
    C_AZUL_HDR    = RGBColor(0x44,0x72,0xC4)
    C_BRANCO      = RGBColor(0xFF,0xFF,0xFF)

# ── Medidas de página (DXA) — baseadas no modelo Toledo de referência ─────────
# Margens: top/bot=2.54cm=1440DXA, left/right=1.91cm=1080DXA
PG_W      = 12242; PG_H    = 15842
MG_TOP    = 1440;  MG_BOT  = 1440
MG_LEFT   = 1080;  MG_RIGHT = 1080
MG_HEADER = 567;   MG_FOOTER = 567
CONTENT_W = PG_W - MG_LEFT - MG_RIGHT  # 10082 DXA ≈ 17.8cm

def cm2emu(c): return int(c * 360000)
def dxa2emu(d): return int(d * 914.4)

# ── Helpers ───────────────────────────────────────────────────────────────────
def _spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:spacing')): pPr.remove(s)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    pPr.append(sp)

def _cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _cell_borders(cell, color='C0C8D8', sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')): tcPr.remove(b)
    brd = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),color)
        brd.append(el)
    # tcBorders deve vir ANTES de shd no CT_TcPr
    shd = tcPr.find(qn('w:shd'))
    if shd is not None: shd.addprevious(brd)
    else: tcPr.append(brd)

def _cell_width(cell, w_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for t in tcPr.findall(qn('w:tcW')): tcPr.remove(t)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(w_dxa)); tcW.set(qn('w:type'),'dxa')
    tcPr.insert(0, tcW)

def _cell_margin(cell, val=120):
    tcPr = cell._tc.get_or_add_tcPr()
    for m in tcPr.findall(qn('w:tcMar')): tcPr.remove(m)
    tcMar = OxmlElement('w:tcMar')
    for side in ['top','left','bottom','right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val)); m.set(qn('w:type'),'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def _no_borders_tbl(tbl):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    for b in tblPr.findall(qn('w:tblBorders')): tblPr.remove(b)
    brd = OxmlElement('w:tblBorders')
    for s in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{s}'); el.set(qn('w:val'),'none')
        brd.append(el)
    # tblBorders deve vir após tblW
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None: tblW.addnext(brd)
    else: tblPr.append(brd)

def _tbl_width(tbl, w_dxa):
    # Define largura via atributo tblLayout=fixed para que o Word respeite as tcW
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    # Remover tblW existente (evita erro de schema)
    for t in tblPr.findall(qn('w:tblW')): tblPr.remove(t)
    # Usar tblLayout fixed para respeitar widths das células
    for t in tblPr.findall(qn('w:tblLayout')): tblPr.remove(t)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def _add_h1_shading(para):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:shd')): pPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),'1A3A6B')
    # Ordem schema CT_PPr: pStyle > pBdr > shd > spacing > ind > jc
    # Inserir shd antes de spacing, ind ou jc (o que vier primeiro)
    ref = None
    for candidate in ['w:spacing','w:ind','w:jc','w:rPr']:
        ref = pPr.find(qn(candidate))
        if ref is not None: break
    if ref is not None: ref.addprevious(shd)
    else: pPr.append(shd)

def _add_h2_border(para):
    pPr = para._p.get_or_add_pPr()
    for b in pPr.findall(qn('w:pBdr')): pPr.remove(b)
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');   bot.set(qn('w:color'),'2E75B6')
    pBdr.append(bot)
    # pBdr vem depois de pStyle/numPr mas antes de shd/spacing/ind
    ref = None
    for candidate in ['w:shd','w:spacing','w:ind','w:jc','w:rPr']:
        ref = pPr.find(qn(candidate))
        if ref is not None: break
    if ref is not None: ref.addprevious(pBdr)
    else: pPr.append(pBdr)

def _add_hyperlink(para, text, url):
    r_id = para.part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), r_id)
    r  = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rSt = OxmlElement('w:rStyle'); rSt.set(qn('w:val'),'Hyperlink')
    rPr.append(rSt); r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t)
    hl.append(r); para._p.append(hl)

def _add_toc(doc):
    p = doc.add_paragraph(); _spacing(p, before=0, after=80)
    run = p.add_run()
    for ftype, text in [('begin',None),('separate',None)]:
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), ftype)
        run._r.append(fc)
        if ftype == 'begin':
            instr = OxmlElement('w:instrText')
            instr.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
            instr.text = ' TOC \\o "1-2" \\h \\z \\u '
            run._r.append(instr)
    hint = p.add_run('Clique com botão direito → "Atualizar campo" para gerar o índice.')
    hint.font.size = Pt(10); hint.font.italic = True
    hint.font.color.rgb = C_AZUL_MED
    p2 = doc.add_paragraph(); _spacing(p2, before=0, after=0)
    run2 = p2.add_run()
    fc_end = OxmlElement('w:fldChar'); fc_end.set(qn('w:fldCharType'),'end')
    run2._r.append(fc_end)

def _add_footer(section, doc_title):
    """Rodapé Toledo: nome do doc à esq | Página X de Y à dir — 8pt, linha sep."""
    ftr = section.footer
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)

    def _fld(para, instr, prefix='', suffix=''):
        if prefix:
            r = para.add_run(prefix)
            r.font.size = Pt(8); r.font.name = 'Arial'
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'),'begin')
        it  = OxmlElement('w:instrText')
        it.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
        it.text = ' ' + instr + ' '
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'separate')
        fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'),'end')
        run = para.add_run()
        run._r.append(fc1); run._r.append(it); run._r.append(fc2); run._r.append(fc3)
        run.font.size = Pt(8); run.font.name = 'Arial'
        if suffix:
            r2 = para.add_run(suffix)
            r2.font.size = Pt(8); r2.font.name = 'Arial'

    # Parágrafo único com tab stop central/direita
    p = ftr.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    # Tab stop à direita — DEVE vir antes de spacing no CT_PPr
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'),  'right')
    tab.set(qn('w:pos'),  str(CONTENT_W))
    tabs.append(tab)
    pPr.append(tabs)
    # Spacing DEPOIS de tabs (ordem CT_PPr: tabs > spacing)
    _spacing(p, before=0, after=0)
    # Linha separadora acima do rodapé
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single'); top.set(qn('w:sz'),  '6')
    top.set(qn('w:space'), '1');      top.set(qn('w:color'),'C0C8D8')
    pBdr.append(top); pPr.insert(0, pBdr)

    # Lado esquerdo: nome do documento
    r_left = p.add_run(doc_title[:60] if doc_title else 'Guardian PRO — Descritivo Funcional')
    r_left.font.size = Pt(8); r_left.font.name = 'Arial'
    r_left.font.color.rgb = RGBColor(0x66,0x66,0x66)

    # Tab para direita
    r_tab = p.add_run(); r_tab._r.append(OxmlElement('w:tab'))

    # Lado direito: Página X de Y
    _fld(p, 'PAGE', prefix='Página ')
    _fld(p, 'NUMPAGES', prefix=' de ')

def _page_break(doc):
    p = doc.add_paragraph(); _spacing(p, before=0, after=0)
    r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'),'page')
    r._r.append(br)

def _sep_line(container, color='1A3A6B', sz=6):
    p = container.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    # pBdr DEVE vir antes de spacing no CT_PPr
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz))
    bot.set(qn('w:space'),'1');   bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)
    # Adicionar spacing depois de pBdr
    _spacing(p, before=60, after=0)

# ── Parser HTML → python-docx ─────────────────────────────────────────────────
def _html_to_docx(doc, html):
    from html.parser import HTMLParser

    class Builder(HTMLParser):
        def __init__(self):
            super().__init__()
            self.stack   = []          # (tag, attrs)
            self.para    = None
            self.in_tbl  = False
            self.tbl_rows = []
            self.cur_row  = []
            self.cur_cell = ''
            self.list_lvl = 0
            self.is_ol    = False
            self.ol_n     = 0

        def _tags(self): return [t for t,_ in self.stack]

        def handle_starttag(self, tag, attrs):
            ad = dict(attrs); self.stack.append((tag, ad))
            if tag == 'h1':
                self.para = doc.add_paragraph(style='Heading 1')
                _spacing(self.para, before=120, after=60)
                self.para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif tag == 'h2':
                self.para = doc.add_paragraph(style='Heading 2')
                _spacing(self.para, before=100, after=40)
                self.para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif tag == 'h3':
                self.para = doc.add_paragraph(style='Heading 3')
                _spacing(self.para, before=80, after=30)
            elif tag == 'p':
                self.para = doc.add_paragraph()
                _spacing(self.para, before=60, after=60)
                self.para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif tag in ('ul','ol'):
                self.list_lvl += 1
                self.is_ol = (tag == 'ol')
                if self.is_ol: self.ol_n = 0
            elif tag == 'li':
                self.para = doc.add_paragraph(); _spacing(self.para,30,30)
                pPr = self.para._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'),    str(720*self.list_lvl))
                ind.set(qn('w:hanging'), '360')
                pPr.append(ind)
                br = self.para.add_run()
                if self.is_ol:
                    self.ol_n += 1; br.text = f'{self.ol_n}.\u00a0'
                else:
                    br.text = '\u2022\u00a0'
                br.font.size = Pt(11)
            elif tag == 'table':
                self.in_tbl = True; self.tbl_rows = []
            elif tag == 'tr':
                self.cur_row = []
            elif tag in ('td','th'):
                self.cur_cell = ''
            elif tag == 'br' and self.para:
                self.para.add_run('\n')

        def handle_endtag(self, tag):
            if self.stack and self.stack[-1][0] == tag: self.stack.pop()
            if tag in ('h1','h2','h3','p','li'): self.para = None
            elif tag in ('ul','ol'):
                self.list_lvl = max(0, self.list_lvl-1)
            elif tag in ('td','th'):
                self.cur_row.append(self.cur_cell.strip())
            elif tag == 'tr':
                if self.cur_row: self.tbl_rows.append(self.cur_row)
            elif tag == 'table':
                self.in_tbl = False
                if not self.tbl_rows: return
                nc = max(len(r) for r in self.tbl_rows)
                nr = len(self.tbl_rows)
                cw = CONTENT_W // nc
                tbl = doc.add_table(rows=nr, cols=nc)
                tbl.style = 'Table Grid'
                _tbl_width(tbl, CONTENT_W)
                for ri, row in enumerate(self.tbl_rows):
                    is_h = (ri == 0)
                    for ci, txt in enumerate(row):
                        if ci >= nc: break
                        cell = tbl.rows[ri].cells[ci]
                        fill = ('4472C4' if is_h else
                                ('F2F2F2' if ri%2==0 else 'FFFFFF'))
                        _cell_shading(cell, fill)
                        _cell_borders(cell)
                        _cell_width(cell, cw)
                        _cell_margin(cell)
                        cp = cell.paragraphs[0]
                        _spacing(cp, 60, 60)
                        run = cp.add_run(txt)
                        run.font.name = 'Arial'; run.font.size = Pt(10)
                        if is_h:
                            run.font.bold = True
                            run.font.color.rgb = C_BRANCO
                self.para = None

        def handle_data(self, data):
            tags = self._tags()
            if self.in_tbl and ('td' in tags or 'th' in tags):
                self.cur_cell += data; return
            if not self.para or not data.strip(): return
            b = 'b' in tags or 'strong' in tags
            i = 'i' in tags or 'em' in tags
            h1 = 'h1' in tags; h2 = 'h2' in tags; h3 = 'h3' in tags
            run = self.para.add_run(data)
            if h1:
                run.font.name='Arial Black'; run.font.size=Pt(13)
                run.font.bold=True; run.font.color.rgb=C_BRANCO
            elif h2:
                run.font.name='Arial Black'; run.font.size=Pt(12)
                run.font.color.rgb=C_AZUL_MED
            elif h3:
                run.font.name='Arial'; run.font.size=Pt(11)
                run.font.bold=True; run.font.color.rgb=C_AZUL_ESCURO
            else:
                run.font.name='Arial'; run.font.size=Pt(11)
                run.font.bold=b; run.font.italic=i

    Builder().feed(html)

    # Aplicar border especial nos H2 gerados pelo HTML
    for p in doc.paragraphs:
        sn = (p.style.name or '') if p.style else ''
        if 'Heading 2' in sn: _add_h2_border(p)


# ── Função principal ──────────────────────────────────────────────────────────
def build_docx_pure(data: dict, toledo_logo: bytes, guardian_banner: bytes) -> bytes:
    doc = Document()

    # Página A4
    sec = doc.sections[0]
    sec.page_width    = Emu(dxa2emu(PG_W))
    sec.page_height   = Emu(dxa2emu(PG_H))
    sec.top_margin      = Emu(dxa2emu(MG_TOP))
    sec.bottom_margin   = Emu(dxa2emu(MG_BOT))
    sec.left_margin     = Emu(dxa2emu(MG_LEFT))
    sec.right_margin    = Emu(dxa2emu(MG_RIGHT))
    sec.header_distance = Emu(dxa2emu(MG_HEADER))
    sec.footer_distance = Emu(dxa2emu(MG_FOOTER))
    # Ativar cabeçalho/rodapé diferentes na 1a página
    # Nota: titlePg removido — rodapé aparece em todas as páginas (incluindo capa)

    # Estilos
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    def _set_outline_lvl(style_elem, lvl):
        pPr = style_elem.get_or_add_pPr()
        for o in pPr.findall(qn('w:outlineLvl')): pPr.remove(o)
        ol = OxmlElement('w:outlineLvl')
        ol.set(qn('w:val'), str(lvl))
        pPr.append(ol)

    for h,nm,sz,rgb,lvl in [
        ('Heading 1','Arial Black',13,C_BRANCO,0),
        ('Heading 2','Arial Black',12,C_AZUL_MED,1),
    ]:
        s = doc.styles[h]
        s.font.name=nm; s.font.size=Pt(sz)
        s.font.color.rgb=rgb; s.font.bold=True
        _set_outline_lvl(s._element, lvl)
    try:
        h3s = doc.styles['Heading 3']
    except:
        h3s = doc.styles.add_style('Heading 3',1)
    h3s.font.name='Arial'; h3s.font.size=Pt(11); h3s.font.bold=True
    h3s.font.color.rgb=C_AZUL_ESCURO
    _set_outline_lvl(h3s._element, 2)

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    hdr = sec.header
    for p in list(hdr.paragraphs): p._element.getparent().remove(p._element)

    htbl = hdr.add_table(rows=1, cols=2, width=Emu(dxa2emu(CONTENT_W)))
    _no_borders_tbl(htbl)
    # Definir largura via layout fixed (sem tblW para evitar erro de schema)
    _tbl_width(htbl, CONTENT_W)
    hw_left  = int(CONTENT_W * 0.60)
    hw_right = CONTENT_W - hw_left
    cl = htbl.rows[0].cells[0]; cr = htbl.rows[0].cells[1]
    _cell_width(cl, hw_left); _cell_width(cr, hw_right)

    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    client_logo_b64 = data.get('clientLogoB64','')
    if client_logo_b64:
        try:
            b64data = client_logo_b64.split('base64,',1)[-1]
            cimg = base64.b64decode(b64data)
            pl.add_run().add_picture(io.BytesIO(cimg), width=Cm(3.0))
        except:
            r=pl.add_run('(logo do cliente)'); r.font.size=Pt(9)
            r.font.italic=True; r.font.color.rgb=RGBColor(0xAA,0xAA,0xAA)
    else:
        r=pl.add_run('(logo do cliente)'); r.font.size=Pt(9)
        r.font.italic=True; r.font.color.rgb=RGBColor(0xAA,0xAA,0xAA)

    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Forçar alinhamento direita via XML (garante no Word)
    pPr_r = pr._p.get_or_add_pPr()
    jc_r  = OxmlElement('w:jc')
    jc_r.set(qn('w:val'), 'right')
    pPr_r.append(jc_r)
    # Usar logo PRIX com fundo branco (substitui logo Toledo)
    try:
        prix_bytes = base64.b64decode(PRIX_LOGO_B64)
        pr.add_run().add_picture(io.BytesIO(prix_bytes), width=Cm(3.5))
    except Exception as e:
        print(f'  [AVISO] Logo PRIX falhou: {e}')
        if toledo_logo:
            pr.add_run().add_picture(io.BytesIO(toledo_logo), width=Cm(3.5))

    # Linha separadora no header (sem spacing para evitar erro de schema)
    p_sep = hdr.add_paragraph()
    pPr_s = p_sep._p.get_or_add_pPr()
    pBdr_s = OxmlElement('w:pBdr')
    bot_s  = OxmlElement('w:bottom')
    bot_s.set(qn('w:val'),'single'); bot_s.set(qn('w:sz'),'6')
    bot_s.set(qn('w:space'),'1');   bot_s.set(qn('w:color'),'1A3A6B')
    pBdr_s.append(bot_s)
    first_s = list(pPr_s)[0] if list(pPr_s) else None
    if first_s is not None: first_s.addprevious(pBdr_s)
    else: pPr_s.append(pBdr_s)

    # ── CAPA ──────────────────────────────────────────────────────────────────
    # Banner Guardian
    if guardian_banner:
        p = doc.add_paragraph(); _spacing(p, 0, 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(guardian_banner),
                                 width=Cm(18.43), height=Cm(7.72))

    # Título DESCRITIVO FUNCIONAL
    p = doc.add_paragraph(); _spacing(p, 120, 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('DESCRITIVO FUNCIONAL')
    r.font.name='Cambria'; r.font.size=Pt(20)
    r.font.bold=True; r.font.color.rgb=C_AZUL_ESCURO

    # Subtítulo
    p = doc.add_paragraph(); _spacing(p, 0, 100)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('GUARDIAN PRO — Software para Gerenciamento de Operações de Pesagem')
    r.font.name='Cambria'; r.font.size=Pt(13)
    r.font.color.rgb=RGBColor(0x22,0x22,0x22)

    # Nome do cliente em destaque
    cn = (data.get('clientName') or '').upper()
    cc = (data.get('clientCity') or '').upper()
    if cn:
        p = doc.add_paragraph(); _spacing(p, 60, 20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cn + (' — ' + cc if cc else ''))
        r.font.name='Cambria'; r.font.size=Pt(16)
        r.font.bold=True; r.font.color.rgb=C_AZUL_ESCURO

    cu = data.get('clientUnit','')
    if cu:
        p = doc.add_paragraph(); _spacing(p, 0, 20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cu)
        r.font.name='Cambria'; r.font.size=Pt(13)
        r.font.color.rgb=RGBColor(0x44,0x44,0x44)

    # Foto da unidade do cliente (se fornecida)
    cib64 = data.get('clientImgB64','')
    if cib64:
        try:
            cimg = base64.b64decode(cib64.split('base64,',1)[-1])
            p = doc.add_paragraph(); _spacing(p, 40, 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(cimg), width=Cm(15))
        except: pass

    # Tabela identificação — layout do modelo: FAZENDA | CT/OV | CT CLOUD
    ct_hw    = data.get('ctHardware','')
    ct_cl    = data.get('ctCloud','')
    filial   = data.get('clientFilial','')
    segmento = data.get('clientSegmento','')

    p_sp = doc.add_paragraph(); _spacing(p_sp, 60, 20)
    W3 = CONTENT_W // 3
    tbl_id = doc.add_table(rows=2, cols=3)
    tbl_id.style = 'Table Grid'
    tbl_id.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_width(tbl_id, CONTENT_W)
    # Cabeçalho da tabela
    hdrs = ['FAZENDA / UNIDADE', 'CT/OV HARDWARE E SERVIÇOS', 'CT CLOUD']
    vals = [cn + (' / ' + cu if cu else ''), ct_hw or '—', ct_cl or '—']
    for ri, (row_data, bg) in enumerate([(hdrs,'1A3A6B'),(vals,'FFFFFF')]):
        for ci, (cell, txt) in enumerate(zip(tbl_id.rows[ri].cells, row_data)):
            _cell_shading(cell, bg)
            _cell_borders(cell); _cell_width(cell, W3); _cell_margin(cell)
            cp = cell.paragraphs[0]; _spacing(cp, 80, 80)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = cp.add_run(txt)
            rr.font.name='Cambria'; rr.font.size=Pt(10)
            rr.font.bold=(ri==0)
            rr.font.color.rgb=C_BRANCO if ri==0 else RGBColor(0x22,0x22,0x22)

    # Filial e segmento
    info_parts = []
    if filial: info_parts.append('FILIAL(IS): ' + filial.upper())
    if segmento: info_parts.append('SEGMENTO: ' + segmento.upper())
    if info_parts:
        p = doc.add_paragraph(); _spacing(p, 30, 20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('   |   '.join(info_parts))
        r.font.name='Cambria'; r.font.size=Pt(10)
        r.font.color.rgb=RGBColor(0x44,0x44,0x44)

    # URL
    p = doc.add_paragraph(); _spacing(p, 40, 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_hyperlink(p, 'www.toledobrasil.com/produto/guardian',
                   'http://www.toledobrasil.com/produto/guardian')

    _sep_line(doc, '1A3A6B', 8)
    _page_break(doc)

    # ── PÁGINA 2: INFORMAÇÕES DO DOCUMENTO + HISTÓRICO ──────────────────────
    def _section_title(doc, txt):
        """Título de seção: texto azul escuro com linha inferior — sem fundo"""
        p = doc.add_paragraph(); _spacing(p, 60, 60)
        r = p.add_run(txt)
        r.font.name='Arial Black'; r.font.size=Pt(13)
        r.font.bold=True; r.font.color.rgb=C_AZUL_ESCURO
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
        bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'1A3A6B')
        pBdr.append(bot); pPr.append(pBdr)

    _section_title(doc, 'Informações do Documento')

    doc_filename = "Descritivo Funcional_GuardianPRO_{}_{}" .format(
        data.get('clientName','NomeCliente'), data.get('docRevision','Rev00'))
    info_rows = [
        ('Título do Documento', 'Descritivo Funcional'),
        ('Autor',               data.get('analystName','') or '—'),
        ('Nome do Arquivo',     doc_filename),
    ]
    TW  = CONTENT_W
    CL2 = int(TW * 0.35)
    CV2 = TW - CL2
    tbl_info = doc.add_table(rows=len(info_rows), cols=2)
    tbl_info.style = 'Table Grid'
    tbl_info.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(tbl_info, TW)
    for ri,(lbl,val) in enumerate(info_rows):
        for cell,w,txt,bold in [
            (tbl_info.rows[ri].cells[0], CL2, lbl, True),
            (tbl_info.rows[ri].cells[1], CV2, val, False),
        ]:
            _cell_shading(cell, 'EBF3FB' if bold else ('FFFFFF' if ri%2==0 else 'F7FAFD'))
            _cell_borders(cell); _cell_width(cell,w); _cell_margin(cell)
            cp = cell.paragraphs[0]; _spacing(cp, 80, 80)
            rr = cp.add_run(txt)
            rr.font.name='Arial'; rr.font.size=Pt(10); rr.font.bold=bold
            if bold: rr.font.color.rgb=C_AZUL_ESCURO

    doc.add_paragraph()
    _section_title(doc, 'Histórico de Revisões')

    rev_data  = data.get('docDate','') or '—'
    rev_num   = data.get('docRevision','Rev00') or 'Rev00'
    rev_desc  = data.get('revDesc','Geração do documento') or 'Geração do documento'
    rev_autor = data.get('analystName','') or '—'

    rev_headers = ['Data', 'Rev.', 'Descrição', 'Autor']
    rev_widths_pct = [0.15, 0.10, 0.50, 0.25]
    rev_widths_dxa = [int(TW * pc) for pc in rev_widths_pct]
    rev_widths_dxa[-1] = TW - sum(rev_widths_dxa[:-1])

    tbl_rev = doc.add_table(rows=2, cols=4)
    tbl_rev.style = 'Table Grid'
    tbl_rev.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width(tbl_rev, TW)
    for ci,(cell,hdr,w) in enumerate(zip(tbl_rev.rows[0].cells, rev_headers, rev_widths_dxa)):
        _cell_shading(cell,'1A3A6B'); _cell_borders(cell); _cell_width(cell,w); _cell_margin(cell)
        cp=cell.paragraphs[0]; _spacing(cp,80,80)
        rr=cp.add_run(hdr); rr.font.name='Arial'; rr.font.size=Pt(10)
        rr.font.bold=True; rr.font.color.rgb=C_BRANCO
    for ci,(cell,val,w) in enumerate(zip(tbl_rev.rows[1].cells,
                                          [rev_data,rev_num,rev_desc,rev_autor], rev_widths_dxa)):
        _cell_shading(cell,'FFFFFF'); _cell_borders(cell); _cell_width(cell,w); _cell_margin(cell)
        cp=cell.paragraphs[0]; _spacing(cp,80,80)
        rr=cp.add_run(str(val)); rr.font.name='Arial'; rr.font.size=Pt(10)

    _sep_line(doc, '1A3A6B', 6)
    _page_break(doc)

    # ── ÍNDICE — título simples, sem fundo colorido ────────────────────────────
    p = doc.add_paragraph(); _spacing(p, 0, 60)
    r = p.add_run('Índice')
    r.font.name='Arial Black'; r.font.size=Pt(14)
    r.font.bold=True; r.font.color.rgb=C_AZUL_ESCURO
    _add_toc(doc)
    _page_break(doc)


    # ── CONTEÚDO ──────────────────────────────────────────────────────────────
    html = data.get('htmlContent','')
    if html:
        _html_to_docx(doc, html)

    # ── IMAGENS PADRÃO GUARDIAN (inseridas após o conteúdo HTML) ──────────────
    def _add_img_section(doc, title, b64_str, caption='', width_cm=16.0):
        """Insere imagem padrão com título e legenda"""
        if not b64_str: return
        try:
            img_data = base64.b64decode(b64_str.split('base64,',1)[-1])
        except: return
        # Título da seção de imagem
        if title:
            p_t = doc.add_paragraph(); _spacing(p_t, 80, 20)
            r_t = p_t.add_run(title)
            r_t.font.name='Arial Black'; r_t.font.size=Pt(12)
            r_t.font.bold=True; r_t.font.color.rgb=C_AZUL_ESCURO
            pPr = p_t._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
            bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'1A3A6B')
            pBdr.append(bot); pPr.append(pBdr)
        # Imagem
        p_i = doc.add_paragraph(); _spacing(p_i, 20, 10)
        p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p_i.add_run().add_picture(io.BytesIO(img_data), width=Cm(width_cm))
        except: return
        # Legenda
        if caption:
            p_c = doc.add_paragraph(); _spacing(p_c, 0, 40)
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_c = p_c.add_run(caption)
            r_c.font.name='Arial'; r_c.font.size=Pt(9)
            r_c.font.italic=True; r_c.font.color.rgb=RGBColor(0x66,0x66,0x66)

    # Extrair imagens do HTML do builder (enviadas como GUARDIAN_IMGS)
    guardian_imgs = data.get('guardianImgs', {})
    mods = data.get('mods', {})

    # Sempre incluir: Arquitetura da Solução
    if guardian_imgs.get('arq_solucao'):
        _page_break(doc)
        _add_img_section(doc,
            'Arquitetura Ilustrativa da Solução',
            guardian_imgs['arq_solucao'],
            'Figura: Arquitetura ilustrativa da solução Guardian PRO',
            width_cm=17.0)

    # Tela de Login / Acesso ao Sistema
    if guardian_imgs.get('tela_login'):
        _add_img_section(doc,
            'Acesso ao Sistema — Tela de Login',
            guardian_imgs['tela_login'],
            'Figura: Tela de autenticação do Guardian PRO',
            width_cm=14.0)

    # Tela principal de operação
    if guardian_imgs.get('tela_operacao'):
        _add_img_section(doc,
            'Tela Principal — Operação',
            guardian_imgs['tela_operacao'],
            'Figura: Tela principal de operação do Guardian PRO',
            width_cm=16.0)

    # Pré-Cadastro (se módulo cadastro)
    if guardian_imgs.get('tela_precadastro'):
        _add_img_section(doc,
            'Tela de Pré-Cadastro',
            guardian_imgs['tela_precadastro'],
            'Figura: Tela de pré-cadastro de veículo no Guardian PRO',
            width_cm=16.0)

    # Cadastramento completo
    if guardian_imgs.get('tela_cadastro'):
        _add_img_section(doc,
            'Tela de Cadastramento',
            guardian_imgs['tela_cadastro'],
            'Figura: Tela de cadastramento de veículo e motorista no Guardian PRO',
            width_cm=16.0)

    # TAG / Cartão
    if guardian_imgs.get('tag_cartao'):
        _add_img_section(doc,
            'Modelo do TAG de Identificação',
            guardian_imgs['tag_cartao'],
            'Figura: Modelo padrão do TAG de identificação veicular — frente e verso',
            width_cm=14.0)

    # Pesagem
    if guardian_imgs.get('tela_pesagem'):
        _add_img_section(doc,
            'Tela de Pesagem',
            guardian_imgs['tela_pesagem'],
            'Figura: Tela de operação de pesagem no Guardian PRO',
            width_cm=16.0)

    # Filas (se módulo fila ativo)
    if guardian_imgs.get('tela_filas_tv') and mods.get('fila'):
        _add_img_section(doc,
            'Gerenciamento de Filas — Exibição em TV',
            guardian_imgs['tela_filas_tv'],
            'Figura: Tela de gerenciamento de filas exibida em painel de TV',
            width_cm=16.0)

    # Inspeção veicular (se módulo inspeção ativo)
    if guardian_imgs.get('tela_inspecao') and mods.get('inspecao'):
        _add_img_section(doc,
            'Tela de Inspeção Veicular',
            guardian_imgs['tela_inspecao'],
            'Figura: Tela de inspeção veicular via dispositivo móvel',
            width_cm=14.0)

    # YMS / Pátios (se módulo YMS ativo)
    if guardian_imgs.get('tela_patio_yms') and mods.get('yms'):
        _add_img_section(doc,
            'Gestão de Pátios (YMS)',
            guardian_imgs['tela_patio_yms'],
            'Figura: Tela de monitoramento de pátios via Cloud Prix (YMS)',
            width_cm=16.0)

    # Rodapé Toledo: nome do doc à esq | Página X de Y à dir
    doc_title = f"Descritivo Funcional_GuardianPRO_{data.get('clientName','')}_"                 f"{data.get('docRevision','Rev00')}"
    _add_footer(sec, doc_title)

    buf = io.BytesIO()
    doc.save(buf)
    # Pós-processamento: corrigir settings.xml e reordenar XML
    import zipfile as _zf, re as _re
    from lxml import etree as _et
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _fix_tbl_order(xml_bytes):
        """Reordena tblPr para garantir tblStyle > jc > tblW > tblBorders > tblLook"""
        root = _et.fromstring(xml_bytes)
        tbl_order = ['tblStyle','tblpPr','jc','tblW','tblBorders',
                     'tblCellSpacing','tblInd','tblCellMar','tblLayout','tblLook']
        for tblPr in root.iter(f'{{{_W}}}tblPr'):
            children = list(tblPr)
            if not children: continue
            # Reordenar conforme tbl_order
            ordered, rest = [], []
            for key in tbl_order:
                for c in children:
                    if c.tag == f'{{{_W}}}{key}':
                        ordered.append(c); break
            rest = [c for c in children if c not in ordered]
            for c in children: tblPr.remove(c)
            for c in ordered + rest: tblPr.append(c)
        return _et.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    buf_in = io.BytesIO(buf.getvalue()); buf_out = io.BytesIO()
    with _zf.ZipFile(buf_in,'r') as zi, _zf.ZipFile(buf_out,'w',_zf.ZIP_DEFLATED) as zo:
        for item in zi.infolist():
            data = zi.read(item.filename)
            if item.filename == 'word/settings.xml':
                xml = data.decode('utf-8')
                xml = _re.sub(r'<w:zoom([^>]*)/>', 
                    lambda m: f'<w:zoom{m.group(1)} w:percent="100"/>'
                    if 'percent' not in m.group(1) else m.group(0), xml)
                data = xml.encode('utf-8')
            elif item.filename in ('word/document.xml','word/header1.xml'):
                try: data = _fix_tbl_order(data)
                except: pass
            zo.writestr(item, data)
    return buf_out.getvalue()

# ══════════════════════════════════════════════════════════════════════

# ── Caminhos das imagens (embutidas no guardian_server.py) ────
# As imagens serão extraídas do builder-descritivo.html
import base64, struct

def extract_imgs_from_builder():
    """Extrai as imagens base64 do builder HTML"""
    imgs = {}
    try:
        html = BUILDER_HTML.read_text(encoding='utf-8')
        for key in ['guardian_capa','logo_toledo_real']:
            m = re.search(rf"{key}:\s*'(data:image/[^']+)'", html)
            if m:
                b64 = m.group(1).split(',',1)[1]
                imgs[key] = base64.b64decode(b64)
    except Exception as e:
        print(f"  [AVISO] Não foi possível extrair imagens: {e}")
    return imgs

# ── Cores Toledo ──────────────────────────────────────────────
C_AZUL     = RGBColor(0x1A,0x3A,0x6B)
C_AZUL_MED = RGBColor(0x2E,0x75,0xB6)
C_AZUL_CLR = RGBColor(0x44,0x72,0xC4)
C_BRANCO   = RGBColor(0xFF,0xFF,0xFF)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def remove_table_borders(tbl_el):
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl_el.insert(0,tblPr)
    for b in tblPr.findall(qn('w:tblBorders')): tblPr.remove(b)
    tblBrd = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto')
        tblBrd.append(el)
    tblPr.append(tblBrd)

def set_para_shd(para, hex_color):
    """shd deve vir antes de spacing/jc no pPr"""
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:shd')): pPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color)
    # Inserir antes de spacing ou jc para manter ordem do schema
    for anchor_tag in [qn('w:spacing'), qn('w:jc'), qn('w:rPr')]:
        anchor = pPr.find(anchor_tag)
        if anchor is not None:
            anchor.addprevious(shd); return
    pPr.append(shd)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    set_para_shd(p, '1A3A6B')
    r = p.add_run(text)
    r.font.name='Arial Black'; r.font.size=Pt(13)
    r.font.bold=True; r.font.color.rgb=C_BRANCO
    # Padding via left indent
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '120')
    pPr.append(ind)

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shd(p, 'DCE6F1')
    r = p.add_run(text)
    r.font.name='Arial Black'; r.font.size=Pt(12)
    r.font.bold=True; r.font.color.rgb=C_AZUL_ESCURO
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'1A3A6B')
    pBdr.append(bot); pPr.append(pBdr)

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name='Arial'; r.font.size=Pt(11)
    r.font.bold=True; r.font.color.rgb=C_AZUL_MED

def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name='Arial'; r.font.size=Pt(11)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name='Arial'; r.font.size=Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)

def make_table(doc, headers, rows, col_widths_cm):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Cabeçalho
    hrow = tbl.rows[0]
    for i,(cell,h) in enumerate(zip(hrow.cells, headers)):
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell,'1A3A6B')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(str(h))
        r.font.name='Arial'; r.font.size=Pt(10)
        r.font.bold=True; r.font.color.rgb=C_BRANCO
    # Dados
    for ri,row in enumerate(rows):
        tr = tbl.rows[ri+1]
        for ci,(cell,val) in enumerate(zip(tr.cells,row)):
            cell.width = Cm(col_widths_cm[ci] if ci<len(col_widths_cm) else 3)
            set_cell_bg(cell,'FFFFFF' if ri%2==0 else 'F2F2F2')
            p = cell.paragraphs[0]
            r = p.add_run(str(val or '—'))
            r.font.name='Arial'; r.font.size=Pt(10)

def strip_html(html):
    if not html: return ''
    h = re.sub(r'<b>(.*?)</b>',r'\1',html,flags=re.I)
    h = re.sub(r'<strong>(.*?)</strong>',r'\1',h,flags=re.I)
    h = re.sub(r'<[^>]+>','',h)
    h = h.replace('&lt;','<').replace('&gt;','>').replace('&amp;','&').replace('&nbsp;',' ')
    return re.sub(r'\s+',' ',h).strip()

def parse_html_content(doc, html):
    """Converte o HTML gerado pelo builder em parágrafos docx"""
    if not html: return
    # Remover CSS, scripts, capa
    html = re.sub(r'<style[\s\S]*?</style>','',html,flags=re.I)
    html = re.sub(r'<script[\s\S]*?</script>','',html,flags=re.I)
    html = re.sub(r'<img[^>]*>','',html,flags=re.I)
    # Pular a capa (até o primeiro H1 depois do page break)
    h1_match = re.search(r'<h1[^>]*>',html,re.I)
    if h1_match: html = html[h1_match.start():]

    for h1part in re.split(r'<h1[^>]*>',html,flags=re.I):
        h1m = re.match(r'(.*?)</h1>',h1part,re.I)
        if not h1m: continue
        add_h1(doc, strip_html(h1m.group(1)))
        rest = h1part[h1m.end():]
        for h2part in re.split(r'<h2[^>]*>',rest,flags=re.I):
            h2m = re.match(r'(.*?)</h2>',h2part,re.I)
            if h2m:
                add_h2(doc, strip_html(h2m.group(1)))
                _parse_block(doc, h2part[h2m.end():])
            else:
                _parse_block(doc, h2part)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _parse_block(doc, html):
    if not html or not html.strip(): return
    tables = re.findall(r'<table[\s\S]*?</table>',html,re.I)
    no_tbl = re.sub(r'<table[\s\S]*?</table>','\n§TBL§\n',html,flags=re.I)
    ti = 0
    for line in no_tbl.split('\n'):
        cl = line.strip()
        if not cl: continue
        if cl == '§TBL§':
            if ti < len(tables): _parse_table(doc, tables[ti]); ti+=1
        elif re.search(r'<h3',cl,re.I):
            t=strip_html(cl)
            if t: add_h3(doc,t)
        elif re.search(r'<li',cl,re.I):
            for item in re.findall(r'<li[^>]*>([\s\S]*?)</li>',cl,re.I) or [cl]:
                t=strip_html(item)
                if t: add_bullet(doc,t)
        else:
            t=strip_html(cl)
            if t and len(t)>2: add_body(doc,t)

def _parse_table(doc, html):
    hdr = re.search(r'<thead[\s\S]*?</thead>',html,re.I)
    headers = [strip_html(h) for h in re.findall(r'<th[^>]*>([\s\S]*?)</th>',
               hdr.group() if hdr else '',re.I)]
    tbody = re.search(r'<tbody[\s\S]*?</tbody>',html,re.I)
    rows = [[strip_html(td) for td in re.findall(r'<td[^>]*>([\s\S]*?)</td>',tr,re.I)]
            for tr in re.findall(r'<tr[^>]*>([\s\S]*?)</tr>',
            tbody.group() if tbody else html,re.I)
            if not re.search(r'<th',tr,re.I)]
    if not headers and not rows: return
    cols = max(len(headers), max((len(r) for r in rows),default=1))
    cw = round(15.27/cols, 2)
    col_widths = [cw]*cols
    make_table(doc, headers or rows[0], rows if headers else rows[1:], col_widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Cache de imagens (carregado uma vez ao iniciar) ───────────
_IMGS_CACHE = {}

def get_imgs():
    global _IMGS_CACHE
    if not _IMGS_CACHE:
        _IMGS_CACHE = extract_imgs_from_builder()
    return _IMGS_CACHE

def b64_src(key):
    """Retorna data URI base64 da imagem"""
    imgs = get_imgs()
    if key not in imgs: return ''
    import base64
    data = base64.b64encode(imgs[key]).decode()
    # Detectar tipo
    magic = imgs[key][:4]
    mime = 'image/png' if magic[:1] == b'\x89' else 'image/jpeg'
    return f"data:{mime};base64,{data}"

def build_html_doc(d):
    """Gera o HTML completo do descritivo — idêntico ao buildDoc() do browser"""
    guardian_src   = b64_src('guardian_capa')
    logo_src       = b64_src('logo_toledo_real')
    client_logo    = d.get('clientLogoB64','')
    client_img     = d.get('clientImgB64','')
    
    client_name = (d.get('clientName') or '').upper()
    client_city = (d.get('clientCity') or '').upper()
    client_unit = d.get('clientUnit','')
    
    # Tabela de identificação
    def row(label, value, even=False):
        bg = '#f4f6fb' if even else '#ffffff'
        return f'''<tr>
          <td style="padding:6pt 9pt;border:.5pt solid #c0c8d8;background:#edf0f7;font-weight:bold;width:150pt;font-family:Arial;font-size:10pt">{label}</td>
          <td style="padding:6pt 9pt;border:.5pt solid #c0c8d8;background:{bg};font-family:Arial;font-size:10pt">{value or 'A definir'}</td>
        </tr>'''

    id_rows = [
        row('CT / OV Hardware e Serviços', d.get('ctHardware'), False),
        row('Licenciamento Cloud', d.get('ctCloud'), True),
    ]
    if d.get('clientFilial'):
        id_rows.append(row('Filial(is)', d['clientFilial'], len(id_rows)%2==0))
    id_rows.append(row('Analista Responsável', d.get('analystName'), len(id_rows)%2==0))
    id_rows.append(row('Revisão', d.get('docRevision','Rev00'), len(id_rows)%2==0))
    id_rows.append(row('Data', d.get('docDate',''), len(id_rows)%2==0))

    # HTML do conteúdo (seções do descritivo)
    html_content = d.get('htmlContent','')

    # Variáveis auxiliares para evitar backslash em f-string (Python 3.11 compat)
    _banner_html = ("<img class='cov-banner' src='" + guardian_src + "' alt='Guardian PRO'>"
                    if guardian_src else "")
    _unit_html   = ("<div class='cov-unit'>" + client_unit + "</div>"
                    if client_unit else "")
    _cimg_html   = ("<img class='cov-client-img' src='" + client_img + "' alt='Unidade'>"
                    if client_img else "")

    html = f"""<!DOCTYPE html>
<html lang="pt-BR"><head><meta charset="UTF-8">
<style>
body{{font-family:Arial,sans-serif;font-size:11pt;color:#1a1a1a;line-height:1.65;margin:0}}
.pg{{max-width:960px;margin:0 auto;padding:40pt}}
h1{{font-family:"Arial Black",Arial;font-size:14pt;color:#fff;background:#1a3a6b;padding:8pt 10pt;margin:24pt 0 10pt;page-break-after:avoid}}
h2{{font-family:"Arial Black",Arial;font-size:12pt;color:#2e75b6;border-bottom:.5pt solid #2e75b6;padding-bottom:3pt;margin:16pt 0 7pt;page-break-after:avoid}}
h3{{font-family:Arial;font-size:11pt;font-weight:bold;color:#2c3e6b;margin:12pt 0 5pt}}
p{{margin:5pt 0;text-align:justify}}
ul,ol{{margin:6pt 0;padding-left:22pt}}
li{{margin-bottom:3pt}}
table{{width:100%;border-collapse:collapse;font-size:10pt;margin:10pt 0;page-break-inside:avoid}}
th{{background:#4472C4;color:#fff;padding:7pt 9pt;text-align:left;font-family:Arial;font-size:9.5pt}}
td{{padding:6pt 9pt;border:.5pt solid #c0c8d8;vertical-align:top;font-family:Arial}}
tr:nth-child(even) td{{background:#f4f6fb}}
.cover{{text-align:center;padding:0 0 24pt;border-bottom:2pt solid #1a3a6b;margin-bottom:24pt}}
.topbar{{display:flex;justify-content:space-between;align-items:center;padding:4pt 0 8pt;border-bottom:1pt solid #c0cce0;margin-bottom:0}}
.cov-logo-client{{height:44pt;width:auto;max-width:200pt;object-fit:contain}}
.cov-logo-placeholder{{color:#aaa;font-style:italic;font-size:9pt}}
.cov-logo-toledo{{height:44pt;width:auto;object-fit:contain}}
.cov-banner{{width:100%;height:auto;display:block;margin:8pt 0 16pt}}
.cov-title{{font-family:"Arial Black",Arial;font-size:22pt;font-weight:bold;color:#1A3A6B;margin-bottom:6pt}}
.cov-sub{{font-size:11pt;color:#444;margin-bottom:18pt}}
.cov-client{{font-family:"Arial Black",Arial;font-size:16pt;font-weight:bold;color:#1A3A6B;margin-bottom:3pt}}
.cov-unit{{font-size:11pt;color:#666;margin-bottom:14pt}}
.cov-client-img{{width:100%;max-height:160pt;object-fit:cover;border-radius:4pt;margin:12pt 0}}
.pbk{{page-break-after:always}}
a{{color:#1a3a6b}}
.note{{background:#fff8e1;border-left:3pt solid #f5a623;padding:8pt 12pt;margin:10pt 0;font-size:10pt}}
.obs{{background:#e8f4fd;border-left:3pt solid #2980b9;padding:8pt 12pt;margin:10pt 0;font-size:10pt}}
</style></head><body><div class="pg">

<div class="cover">
    {_banner_html}
  <div class="cov-title">DESCRITIVO FUNCIONAL</div>
  <div class="cov-sub">GUARDIAN PRO — Software para Gerenciamento de Operações de Pesagem</div>
  <div class="cov-client">{client_name}{" — " + client_city if client_city else ""}</div>
  {_unit_html}
  {_cimg_html}
  <table style="width:auto;margin:14pt auto"><tbody>{''.join(id_rows)}</tbody></table>
  <p style="font-size:9pt;color:#888"><a href="http://www.toledobrasil.com/produto/guardian">www.toledobrasil.com/produto/guardian</a></p>
</div>
<div class="pbk"></div>

{html_content}

</div></body></html>"""
    return html

# ── Pós-processamento: formatação Toledo no DOCX gerado pelo LibreOffice ──
# Usa lxml + addprevious() para inserir elementos na posição EXATA do schema OOXML

from lxml import etree as _ET

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Elementos que vêm APÓS shd no schema CT_PPr
_AFTER_SHD = {'tabs','suppressAutoHyphens','kinsoku','wordWrap','overflowPunct',
    'topLinePunct','autoSpaceDE','autoSpaceDN','bidi','adjustRightInd','snapToGrid',
    'spacing','ind','contextualSpacing','mirrorIndents','suppressOverlap','jc',
    'textDirection','textAlignment','textboxTightWrap','outlineLvl','divId',
    'cnfStyle','rPr','sectPr','pPrChange'}
_AFTER_PBDR = {'shd'} | _AFTER_SHD

# Elementos que vêm APÓS shd no schema CT_TcPr
_AFTER_TCPR_SHD = {'noWrap','tcMar','textDirection','tcFitText','vAlign',
                   'hideMark','headers','tcPrChange'}

def _first_child_in(parent, tag_set):
    for child in parent:
        if not isinstance(child.tag, str): continue
        if _ET.QName(child.tag).localname in tag_set:
            return child
    return None

def _insert_ppr(ppr, new_elem, after_set):
    for e in ppr.findall(new_elem.tag): ppr.remove(e)
    anchor = _first_child_in(ppr, after_set)
    if anchor is not None: anchor.addprevious(new_elem)
    else: ppr.append(new_elem)

def _insert_tcpr_shd(tcpr, hex_color):
    new_shd = _ET.Element(f'{{{_W}}}shd')
    new_shd.set(f'{{{_W}}}val','clear')
    new_shd.set(f'{{{_W}}}color','auto')
    new_shd.set(f'{{{_W}}}fill', hex_color)
    for e in tcpr.findall(f'{{{_W}}}shd'): tcpr.remove(e)
    anchor = _first_child_in(tcpr, _AFTER_TCPR_SHD)
    if anchor is not None: anchor.addprevious(new_shd)
    else: tcpr.append(new_shd)

def _make_shd(hex_color):
    s = _ET.Element(f'{{{_W}}}shd')
    s.set(f'{{{_W}}}val','clear'); s.set(f'{{{_W}}}color','auto')
    s.set(f'{{{_W}}}fill', hex_color)
    return s

def _make_pbdr(color='2E75B6'):
    p = _ET.Element(f'{{{_W}}}pBdr')
    b = _ET.SubElement(p, f'{{{_W}}}bottom')
    b.set(f'{{{_W}}}val','single'); b.set(f'{{{_W}}}sz','4')
    b.set(f'{{{_W}}}space','1'); b.set(f'{{{_W}}}color', color)
    return p

def _apply_toledo_formatting(docx_path, logo_client_b64=""):
    from docx import Document as _Doc
    from docx.shared import Pt as _Pt, RGBColor as _RGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    import io as _io

    doc = _Doc(str(docx_path))

    for para in doc.paragraphs:
        sname = (para.style.name or '') if para.style else ''
        sl = sname.lower()
        ppr = para._p.find(f'{{{_W}}}pPr')
        if ppr is None: continue

        if 'heading 1' in sl or sname in {'Heading1','Ttulo1'}:
            _insert_ppr(ppr, _make_shd('1A3A6B'), _AFTER_SHD)
            para.alignment = _AL.LEFT
            para.paragraph_format.space_before = _Pt(10)
            para.paragraph_format.space_after  = _Pt(5)
            for r in para.runs:
                r.font.color.rgb = _RGB(0xFF,0xFF,0xFF)
                r.font.name = 'Arial Black'
                r.font.size = _Pt(13)
                r.font.bold = True

        elif 'heading 2' in sl or sname in {'Heading2','Ttulo2'}:
            _insert_ppr(ppr, _make_pbdr(), _AFTER_PBDR)
            para.alignment = _AL.LEFT
            para.paragraph_format.space_before = _Pt(8)
            para.paragraph_format.space_after  = _Pt(4)
            for r in para.runs:
                r.font.color.rgb = _RGB(0x2E,0x75,0xB6)
                r.font.name = 'Arial Black'
                r.font.size = _Pt(12)

        else:
            if para.alignment is None:
                para.alignment = _AL.JUSTIFY
            for r in para.runs:
                if not r.font.name: r.font.name = 'Arial'
                if not r.font.size: r.font.size = _Pt(11)

    # Encontrar posição do primeiro H1 para identificar tabela da capa
    first_h1_idx = None
    for pi, para in enumerate(doc.paragraphs):
        if 'heading 1' in ((para.style.name or '') if para.style else '').lower():
            first_h1_idx = pi; break

    for ti, table in enumerate(doc.tables):
        # Detectar se é tabela da capa: vem antes do primeiro H1
        # Localizar posição da tabela no documento
        tbl_elem = table._tbl
        is_cover_table = False
        if first_h1_idx is not None:
            # Verificar se a tabela está antes do primeiro H1 no XML
            body = tbl_elem.getparent()
            if body is not None:
                children = list(body)
                tbl_pos = children.index(tbl_elem) if tbl_elem in children else 999
                # Encontrar posição do primeiro H1
                h1_pos = 999
                for child in children:
                    if child.tag == f'{{{_W}}}p':
                        ppr = child.find(f'{{{_W}}}pPr')
                        if ppr is not None:
                            ps = ppr.find(f'{{{_W}}}pStyle')
                            if ps is not None and 'Heading1' in ps.get(f'{{{_W}}}val',''):
                                h1_pos = children.index(child); break
                is_cover_table = tbl_pos < h1_pos

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if is_cover_table:
                    # Tabela da capa: col esq = label (cinza), col dir = valor (branco)
                    color = 'EDF0F7' if ci == 0 else ('FFFFFF' if ri%2==0 else 'F4F6FB')
                    is_hdr = False
                else:
                    is_hdr = ri == 0
                    color = '4472C4' if is_hdr else ('FFFFFF' if ri%2==1 else 'F2F2F2')
                tcpr = cell._tc.find(f'{{{_W}}}tcPr')
                if tcpr is None:
                    tcpr = _ET.SubElement(cell._tc, f'{{{_W}}}tcPr')
                _insert_tcpr_shd(tcpr, color)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = _Pt(10)
                        if is_hdr:
                            r.font.color.rgb = _RGB(0xFF,0xFF,0xFF)
                            r.font.bold = True
                        elif is_cover_table and ci == 0:
                            r.font.bold = True

    buf = _io.BytesIO()
    doc.save(buf)
    return _fix_settings(buf.getvalue(), logo_client_b64=logo_client_b64)

def _make_header_xml(logo_client_b64, logo_toledo_b64, rId_client, rId_toledo):
    """Gera o XML do cabeçalho com logos cliente (esq) e Toledo (dir)"""
    WNS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    RNS = 'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
    WPNS = 'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    ANS = 'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    PICNS = 'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'

    def inline_img(rId, cx, cy, desc, img_id):
        return f"""<wp:inline distT="0" distB="0" distL="0" distR="0">
          <wp:extent cx="{cx}" cy="{cy}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:docPr id="{img_id}" name="{desc}"/>
          <wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:nvPicPr>
                  <pic:cNvPr id="{img_id}" name="{desc}"/>
                  <pic:cNvPicPr><a:picLocks noChangeAspect="1" noChangeArrowheads="1"/></pic:cNvPicPr>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="{rId}"/>
                  <a:stretch><a:fillRect/></a:stretch>
                </pic:blipFill>
                <pic:spPr bwMode="auto">
                  <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:noFill/>
                  <a:ln><a:noFill/></a:ln>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>
        </wp:inline>"""

    # Toledo logo: 4.86cm x 2.04cm → EMU (1cm = 914400/100 = 9144... 1cm = 360000 EMU)
    # 1 cm = 360000 EMU
    cx_toledo = int(4.86 * 360000)   # 1749600
    cy_toledo = int(2.04 * 360000)   # 734400
    # Logo cliente placeholder: mesmo tamanho
    cx_client = int(3.0 * 360000)    # 1080000
    cy_client = int(2.04 * 360000)   # 734400

    # Se tem logo cliente, usa imagem; senão, texto placeholder
    if logo_client_b64 and rId_client:
        cell_left = f"""<w:p><w:pPr><w:jc w:val="left"/></w:pPr>
        <w:r><w:rPr><w:noProof/></w:rPr>
          <w:drawing>{inline_img(rId_client, cx_client, cy_client, "Logo Cliente", 10)}</w:drawing>
        </w:r></w:p>"""
    else:
        cell_left = """<w:p><w:pPr><w:jc w:val="left"/></w:pPr>
        <w:r><w:rPr>
          <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
          <w:color w:val="AAAAAA"/><w:sz w:val="18"/>
        </w:rPr><w:t>(logo do cliente)</w:t></w:r></w:p>"""

    cell_right = f"""<w:p><w:pPr><w:jc w:val="right"/></w:pPr>
      <w:r><w:rPr><w:noProof/></w:rPr>
        <w:drawing>{inline_img(rId_toledo, cx_toledo, cy_toledo, "Toledo do Brasil", 11)}</w:drawing>
      </w:r></w:p>"""

    # Linha separadora abaixo do cabeçalho
    sep_line = """<w:p><w:pPr>
      <w:pBdr><w:bottom w:val="single" w:sz="6" w:space="1" w:color="C0C8D8"/></w:pBdr>
    </w:pPr></w:p>"""

    xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr {WNS} {RNS} {WPNS} {ANS} {PICNS}
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:tbl>
    <w:tblPr>
      <w:tblW w:w="0" w:type="auto"/>
      <w:tblBorders>
        <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      </w:tblBorders>
    </w:tblPr>
    <w:tblGrid>
      <w:gridCol w:w="4500"/>
      <w:gridCol w:w="4500"/>
    </w:tblGrid>
    <w:tr>
      <w:tc>
        <w:tcPr><w:tcW w:w="4500" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
        {cell_left}
      </w:tc>
      <w:tc>
        <w:tcPr><w:tcW w:w="4500" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
        {cell_right}
      </w:tc>
    </w:tr>
  </w:tbl>
  {sep_line}
</w:hdr>"""
    return xml.encode('utf-8')


def _make_toc_xml():
    """Gera o XML do índice (TOC) que o Word atualiza automaticamente"""
    return """<w:p>
  <w:pPr><w:pStyle w:val="TOCHeading"/><w:spacing w:before="240" w:after="120"/></w:pPr>
  <w:r><w:rPr><w:rFonts w:ascii="Arial Black" w:hAnsi="Arial Black"/><w:b/><w:color w:val="1A3A6B"/><w:sz w:val="28"/></w:rPr>
    <w:t>Índice</w:t>
  </w:r>
</w:p>
<w:p>
  <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>
  <w:fldSimple w:instr=" TOC \\o &quot;1-2&quot; \\h \\z \\u ">
    <w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:sz w:val="20"/><w:color w:val="1A3A6B"/></w:rPr>
      <w:t>Clique com botão direito e &quot;Atualizar campo&quot; para gerar o índice.</w:t>
    </w:r>
  </w:fldSimple>
</w:p>
<w:p><w:pPr><w:pageBreakBefore/><w:spacing w:before="0" w:after="0"/></w:pPr></w:p>"""


def _fix_settings(docx_bytes, logo_client_b64='', logo_toledo_b64=''):
    buf_in=io.BytesIO(docx_bytes); buf_out=io.BytesIO()
    
    with zipfile.ZipFile(buf_in,'r') as zi:
        # Ler todos os arquivos
        all_files = {item.filename: zi.read(item.filename) for item in zi.infolist()}
    
    # ── Preparar imagens para o cabeçalho ──
    imgs_raw = extract_imgs_from_builder()
    toledo_png = imgs_raw.get('logo_toledo_real', b'')
    
    # IDs dos relacionamentos do cabeçalho
    rId_toledo = 'rHdr1'
    rId_client = 'rHdr2' if logo_client_b64 else ''
    
    # ── Gerar header1.xml ──
    header_xml = _make_header_xml(logo_client_b64, logo_toledo_b64, rId_client, rId_toledo)
    all_files['word/header1.xml'] = header_xml
    
    # ── Gerar header1.xml.rels ──
    rels_hdr = [
        f'<Relationship Id="{rId_toledo}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
        f'Target="media/image_toledo_hdr.jpeg"/>',
    ]
    if logo_client_b64 and rId_client:
        rels_hdr.append(
            f'<Relationship Id="{rId_client}" '
            f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            f'Target="media/image_client_hdr.png"/>'
        )
    hdr_rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(rels_hdr) +
        '</Relationships>'
    ).encode('utf-8')
    all_files['word/_rels/header1.xml.rels'] = hdr_rels_xml
    
    # ── Salvar logo Toledo na pasta media do header ──
    if toledo_png:
        all_files['word/media/image_toledo_hdr.jpeg'] = toledo_png
    
    # ── Salvar logo cliente se fornecido ──
    if logo_client_b64:
        try:
            import base64 as _b64
            client_data = _b64.b64decode(logo_client_b64.split(",",1)[-1])
            all_files['word/media/image_client_hdr.png'] = client_data
        except: pass
    
    # ── Atualizar document.xml.rels para incluir header ──
    rels_doc = all_files.get('word/_rels/document.xml.rels', b'').decode('utf-8')
    if 'header1.xml' not in rels_doc:
        rels_doc = rels_doc.replace(
            '</Relationships>',
            '<Relationship Id="rHdrMain" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" '
            'Target="header1.xml"/></Relationships>'
        )
    all_files['word/_rels/document.xml.rels'] = rels_doc.encode('utf-8')
    
    # ── Atualizar document.xml: sectPr + TOC ──
    doc_xml = all_files.get('word/document.xml', b'').decode('utf-8')
    
    # Adicionar headerReference no sectPr
    if 'w:headerReference' not in doc_xml:
        doc_xml = doc_xml.replace(
            '<w:type w:val="nextPage"/>',
            '<w:headerReference w:type="default" r:id="rHdrMain"/>'
            '<w:type w:val="nextPage"/>'
        )
        # Se não tinha <w:type>, tentar outro ponto
        if 'w:headerReference' not in doc_xml:
            doc_xml = re.sub(
                r'(<w:sectPr[^>]*>)',
                r'\1<w:headerReference w:type="default" r:id="rHdrMain"/>',
                doc_xml
            )
    
    # Aumentar margem do topo para cabeçalho (header=851=1.5cm, top=851)
    doc_xml = re.sub(
        r'w:header="[^"]*"',
        'w:header="851"',
        doc_xml
    )
    doc_xml = re.sub(
        r'(w:top=")567(")',
        r'\g<1>1417\g<2>',  # 2.5cm top margin
        doc_xml
    )
    
    # Inserir TOC após o page-break da capa (primeiro <w:p> com pageBreakBefore ou após pbk)
    # O documento tem: capa → <p pageBreakBefore> → conteúdo
    # Inserir TOC + page-break APÓS o primeiro page-break
    toc_xml = _make_toc_xml()
    # Inserir TOC: buscar page-break da capa ou primeiro H1
    pbk_idx = doc_xml.find('<w:pageBreakBefore/>')
    h1_idx  = doc_xml.find('<w:pStyle w:val="Heading1"/>')
    if pbk_idx > 0:
        end_p = doc_xml.find('</w:p>', pbk_idx)
        if end_p > 0:
            insert_pos = end_p + len('</w:p>')
            doc_xml = doc_xml[:insert_pos] + toc_xml + doc_xml[insert_pos:]
    elif h1_idx > 0:
        start_p = doc_xml.rfind('<w:p>', 0, h1_idx)
        if start_p > 0:
            doc_xml = doc_xml[:start_p] + toc_xml + '<w:p><w:pPr><w:pageBreakBefore/></w:pPr></w:p>' + doc_xml[start_p:]
    all_files['word/document.xml'] = doc_xml.encode('utf-8')
    
    # ── Processar demais arquivos (settings, styles) ──
    if 'word/settings.xml' in all_files:
        xml = all_files['word/settings.xml'].decode('utf-8')
        xml = re.sub(r'<w:zoom([^>]*)/>', 
            lambda m: f'<w:zoom{m.group(1)} w:percent="100"/>' if 'percent' not in m.group(1) else m.group(0), xml)
        all_files['word/settings.xml'] = xml.encode('utf-8')
    
    if 'word/styles.xml' in all_files:
        xml = all_files['word/styles.xml'].decode('utf-8')
        xml = xml.replace('w:ascii="Arial Black;Arial"','w:ascii="Arial Black"')
        xml = xml.replace('w:hAnsi="Arial Black;Arial"','w:hAnsi="Arial Black"')
        xml = xml.replace('w:eastAsia="Arial Black;Arial"','w:eastAsia="Arial Black"')
        xml = xml.replace('w:cs="Arial Black;Arial"','w:cs="Arial Black"')
        xml = re.sub(r'(<w:style[^>]*styleId="Heading1".*?<w:pPr>.*?)<w:ind[^/]*/>', r'\1<w:ind w:left="0" w:right="0"/>', xml, flags=re.DOTALL)
        xml = re.sub(r'(<w:style[^>]*styleId="Heading1".*?)<w:pBdr/>', r'\1', xml, flags=re.DOTALL)
        # Adicionar estilo TOCHeading se não existir
        if 'TOCHeading' not in xml:
            toc_style = """<w:style w:type="paragraph" w:styleId="TOCHeading">
  <w:name w:val="TOC Heading"/>
  <w:basedOn w:val="Normal"/>
  <w:pPr><w:spacing w:before="240" w:after="120"/></w:pPr>
  <w:rPr><w:rFonts w:ascii="Arial Black" w:hAnsi="Arial Black"/><w:b/><w:color w:val="1A3A6B"/><w:sz w:val="26"/></w:rPr>
</w:style>"""
            xml = xml.replace('</w:styles>', toc_style + '</w:styles>')
        all_files['word/styles.xml'] = xml.encode('utf-8')
    
    # ── Remontar o ZIP ──
    with zipfile.ZipFile(buf_out,'w',zipfile.ZIP_DEFLATED) as zo:
        for fname, data in all_files.items():
            zo.writestr(fname, data)
    
    return buf_out.getvalue()


# ── GERADOR PRINCIPAL ─────────────────────────────────────────
def generate_docx(data, imgs):
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(1.9)
        section.right_margin  = Cm(1.9)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    # ── CAPA ──
    # Tabela topo: logo cliente (esq) | logo Toledo (dir)
    tbl_top = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl_top._tbl)
    c_l = tbl_top.rows[0].cells[0]; c_l.width = Cm(8.27)
    c_r = tbl_top.rows[0].cells[1]; c_r.width = Cm(9.00)

    # Logo cliente
    if data.get('clientLogoB64'):
        try:
            logo_data = base64.b64decode(data['clientLogoB64'].split(',',1)[1])
            p_l = c_l.paragraphs[0]; p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_l.add_run().add_picture(io.BytesIO(logo_data), height=Cm(2.04))
        except:
            _placeholder_logo(c_l)
    else:
        _placeholder_logo(c_l)

    # Logo Toledo
    p_r = c_r.paragraphs[0]; p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if 'logo_toledo_real' in imgs:
        p_r.add_run().add_picture(io.BytesIO(imgs['logo_toledo_real']),
                                   width=Cm(4.86), height=Cm(2.04))
    else:
        r = p_r.add_run('Toledo do Brasil')
        r.font.name='Arial Black'; r.font.size=Pt(14)
        r.font.bold=True; r.font.color.rgb=C_AZUL_MED

    # Espaço
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Banner Guardian — 18.43×7.72cm, LEFT
    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_banner.paragraph_format.space_before = Pt(2)
    p_banner.paragraph_format.space_after  = Pt(10)
    if 'guardian_capa' in imgs:
        p_banner.add_run().add_picture(io.BytesIO(imgs['guardian_capa']),
                                        width=Cm(18.43), height=Cm(7.72))

    # Título
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('DESCRITIVO FUNCIONAL')
    r.font.name='Arial Black'; r.font.size=Pt(22)
    r.font.bold=True; r.font.color.rgb=C_AZUL

    # Subtítulo
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run('GUARDIAN PRO — Software para Gerenciamento de Operações de Pesagem')
    r.font.name='Arial'; r.font.size=Pt(12)
    r.font.color.rgb=RGBColor(0x44,0x44,0x44)

    # Nome cliente
    client_str = (data.get('clientName','') or '').upper()
    if data.get('clientCity'): client_str += f" — {data['clientCity'].upper()}"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(client_str)
    r.font.name='Arial Black'; r.font.size=Pt(16)
    r.font.bold=True; r.font.color.rgb=C_AZUL

    if data.get('clientUnit'):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(data['clientUnit'])
        r.font.name='Cambria'; r.font.size=Pt(13)
        r.font.color.rgb=RGBColor(0x66,0x66,0x66)

    # Tabela identificação
    rows_id = [
        ('CT / OV Hardware e Serviços', data.get('ctHardware','A definir')),
        ('Licenciamento Cloud',          data.get('ctCloud','A definir')),
        ('Analista Responsável',         data.get('analystName','—')),
        ('Revisão',                      data.get('docRevision','Rev00')),
        ('Data do Documento',            data.get('docDate','—')),
    ]
    if data.get('clientFilial'):
        rows_id.insert(2, ('Filial(is)', data['clientFilial']))
    make_table(doc, ['Campo','Valor'], rows_id, [5.5, 9.77])

    # URL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run('www.toledobrasil.com/produto/guardian')
    r.font.name='Arial'; r.font.size=Pt(9)
    r.font.italic=True; r.font.color.rgb=C_AZUL

    doc.add_page_break()

    # ── HISTÓRICO DE REVISÕES ──
    add_h1(doc, 'Informações do Documento')
    make_table(doc,
        ['Data','Revisão','Descrição','Analista'],
        [[data.get('docDate','—'), data.get('docRevision','Rev00'),
          data.get('revDesc','Geração do documento'), data.get('analystName','—')]],
        [2.5, 1.5, 8.27, 3.0])
    doc.add_page_break()

    # ── CONTEÚDO PRINCIPAL ──
    if data.get('htmlContent'):
        parse_html_content(doc, data['htmlContent'])

    # Salvar
    buf = io.BytesIO()
    doc.save(buf)
    return _fix_settings(buf.getvalue())

def _placeholder_logo(cell):
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('(logo do cliente)')
    r.font.italic=True; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor(0xAA,0xAA,0xAA)


# ── SERVIDOR HTTP ─────────────────────────────────────────────
class Handler(http.server.BaseHTTPRequestHandler):
    def log_message(self, fmt, *args):
        print(f"  [{args[1]}] {args[0]}")

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_POST(self):
        if self.path == '/convert':
            # Recebe MHT do browser, converte para DOCX via LibreOffice
            length = int(self.headers.get('Content-Length', 0))
            mht_data = self.rfile.read(length)
            tmp_dir = out_dir = None
            try:
                tmp_dir = tempfile.mkdtemp()
                mht_path = os.path.join(tmp_dir, 'descritivo.mht')
                with open(mht_path, 'wb') as f:
                    f.write(mht_data)
                out_dir = tempfile.mkdtemp()
                # Configurar ambiente para LibreOffice no container Railway
                lo_home = tempfile.mkdtemp(prefix='lo_home_')
                lo_env = os.environ.copy()
                lo_env['HOME'] = lo_home
                lo_env['TMPDIR'] = lo_home

                result = subprocess.run(
                    ['soffice',
                     '--headless',
                     '--norestore',
                     '--nofirststartwizard',
                     '--nolockcheck',
                     '--convert-to', 'docx',
                     '--outdir', out_dir,
                     mht_path],
                    capture_output=True, text=True,
                    timeout=120, env=lo_env
                )
                # Limpar HOME temporário do LibreOffice
                shutil.rmtree(lo_home, ignore_errors=True)
                docx_files = list(Path(out_dir).glob('*.docx'))
                if not docx_files:
                    raise Exception(f'LibreOffice falhou: {result.stderr[:300]}')
                # Pós-processamento: aplicar formatação Toledo via python-docx
                docx_bytes = _apply_toledo_formatting(docx_files[0], logo_client_b64=data.get('clientLogoB64',''))
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Length', len(docx_bytes))
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(docx_bytes)
                print(f'  ✓ MHT→DOCX: {len(docx_bytes)//1024}KB')
            except subprocess.TimeoutExpired:
                print('  ✗ LibreOffice timeout')
                self._error('Timeout na conversão — tente novamente')
            except Exception as e:
                import traceback
                print(f'  ✗ Erro conversão: {e}')
                traceback.print_exc()
                self._error(str(e))
            finally:
                if tmp_dir: shutil.rmtree(tmp_dir, ignore_errors=True)
                if out_dir: shutil.rmtree(out_dir, ignore_errors=True)

        elif self.path == '/generate':
            # Recebe JSON → gera DOCX via python-docx puro (sem LibreOffice)
            if not _DOCX_AVAILABLE:
                self._error('python-docx não disponível no servidor. Verifique requirements.txt.')
                return
            length = int(self.headers.get('Content-Length', 0))
            body   = self.rfile.read(length)
            try:
                data = json.loads(body)

                # Carregar imagens Toledo do builder HTML
                imgs        = extract_imgs_from_builder()
                toledo_logo = imgs.get('logo_toledo_real', b'')
                guardian_bn = imgs.get('guardian_capa',    b'')

                # Gerar DOCX puro
                docx_bytes = build_docx_pure(data, toledo_logo, guardian_bn)

                self.send_response(200)
                self.send_header('Content-Type',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Length', len(docx_bytes))
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(docx_bytes)
                print(f'  ✓ DOCX puro: {len(docx_bytes)//1024}KB')

                # Salvar no histórico (silencioso em caso de falha)
                try:
                    _salvar_projeto(data)
                except Exception as _dbe:
                    print(f'  [DB] Aviso: {_dbe}')

            except Exception as e:
                import traceback; traceback.print_exc()
                self._error(str(e))

    def do_GET(self):
        from urllib.parse import urlparse, parse_qs
        path = self.path.split('?')[0]
        qs   = parse_qs(urlparse(self.path).query)

        # ── Builder (rota principal) ──────────────────────────────
        if path in ('/', '/builder'):
            try:
                content = BUILDER_HTML.read_bytes()
                self.send_response(200)
                self.send_header('Content-Type', 'text/html; charset=utf-8')
                self.send_header('Content-Length', len(content))
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(content)
            except FileNotFoundError:
                self.send_error(404, 'builder-descritivo.html nao encontrado')

        # ── Histórico: listar ─────────────────────────────────────
        elif path == '/projetos':
            analista = qs.get('analista', [''])[0]
            busca    = qs.get('busca',    [''])[0]
            segmento = qs.get('segmento', [''])[0]
            projetos = _listar_projetos(analista=analista, busca=busca, segmento=segmento)
            body = json.dumps({'projetos': projetos}, ensure_ascii=False).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', len(body))
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(body)

        # ── Histórico: clonar ─────────────────────────────────────
        elif path == '/projetos/clonar':
            proj_id = qs.get('id', [''])[0]
            payload = _carregar_projeto(proj_id) if proj_id else None
            if payload:
                body = json.dumps({'ok': True, 'payload': payload}, ensure_ascii=False).encode()
                self.send_response(200)
            else:
                body = json.dumps({'ok': False, 'erro': 'Projeto nao encontrado'}).encode()
                self.send_response(404)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', len(body))
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(body)

        # ── Histórico: excluir ────────────────────────────────────
        elif path == '/projetos/excluir':
            proj_id = qs.get('id', [''])[0]
            ok = _excluir_projeto(proj_id) if proj_id else False
            body = json.dumps({'ok': ok}).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Content-Length', len(body))
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(body)

        # ── Health check ──────────────────────────────────────────
        elif path == '/ping':
            body = b'pong'
            self.send_response(200)
            self.send_header('Content-Type', 'text/plain')
            self.send_header('Content-Length', len(body))
            self.end_headers()
            self.wfile.write(body)

        # ── Changelog ─────────────────────────────────────────────
        elif path == '/changelog':
            changelog_file = HERE / 'changelog.json'
            if changelog_file.exists():
                try:
                    data = changelog_file.read_bytes()
                    self.send_response(200)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.send_header('Content-Length', len(data))
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                    self.wfile.write(data)
                except Exception as e:
                    self._error(str(e))
            else:
                body = b'[]'
                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Content-Length', len(body))
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(body)

        # ── Status do banco de dados ───────────────────────────────
        elif path == '/db-status':
            db_url = os.environ.get('DATABASE_URL', '')
            if not db_url:
                status = {'ok': False, 'erro': 'DATABASE_URL não configurada. Adicione o PostgreSQL no Railway.'}
            else:
                try:
                    conn = _get_db()
                    if conn:
                        cur = conn.cursor()
                        cur.execute("SELECT COUNT(*) FROM projetos WHERE status='ativo'")
                        total = cur.fetchone()[0]
                        cur.close(); conn.close()
                        status = {'ok': True, 'total_projetos': total, 'banco': 'conectado'}
                    else:
                        status = {'ok': False, 'erro': 'Falha ao conectar no banco'}
                except Exception as e:
                    status = {'ok': False, 'erro': str(e)}
            body = json.dumps(status, ensure_ascii=False).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', len(body))
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(body)

        else:
            self.send_response(404)
            self.end_headers()

    def send_header_cors(self):
        """Helper para não repetir CORS em toda rota."""
        pass  # CORS é adicionado nas respostas individualmente

    def _error(self, msg):
        self.send_response(500)
        data = msg.encode('utf-8')
        self.send_header('Content-Type','text/plain; charset=utf-8')
        self.send_header('Content-Length', len(data))
        self.send_header('Access-Control-Allow-Origin','*')
        self.end_headers()
        self.wfile.write(data)


if __name__ == '__main__':
    _init_db()  # Garante tabela no PostgreSQL
    import socketserver

    print()
    print('  Construtor de Descritivo Funcional — Servidor iniciando...')
    print(f'  Porta: {PORT}')
    print(f'  builder.html: {"OK" if BUILDER_HTML.exists() else "NAO ENCONTRADO"}')
    db_url = os.environ.get('DATABASE_URL', '')
    if db_url:
        print('  DATABASE_URL: configurada ✓')
    else:
        print('  DATABASE_URL: NÃO configurada — histórico desativado')
        print('  → Adicione PostgreSQL no Railway para habilitar o histórico')
    print()

    if not BUILDER_HTML.exists():
        print(f'  [ERRO] builder-descritivo.html nao encontrado em {BUILDER_HTML.parent}')
        sys.exit(1)

    import signal
    def shutdown(signum, frame):
        print('\n  Encerrando servidor...')
        sys.exit(0)
    signal.signal(signal.SIGTERM, shutdown)
    signal.signal(signal.SIGINT, shutdown)

    try:
        socketserver.TCPServer.allow_reuse_address = True
        with socketserver.TCPServer(('0.0.0.0', PORT), Handler) as httpd:
            print(f'  Servidor rodando na porta {PORT}')
            httpd.serve_forever()
    except KeyboardInterrupt:
        print('\n  Servidor encerrado.')
    except OSError as e:
        print(f'\n  [ERRO] Porta {PORT}: {e}')
        sys.exit(1)
    